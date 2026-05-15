from dotenv import load_dotenv
from pathlib import Path
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

import os
import uuid
import hashlib
import logging
import json
import time
from copy import deepcopy
import bcrypt
import jwt
import requests
try:
    import ipfshttpclient  # optional; HTTP API is primary for compatibility
except Exception:  # pragma: no cover
    ipfshttpclient = None
from datetime import datetime, timezone, timedelta
from typing import Optional, List
try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

from fastapi import FastAPI, APIRouter, Request, Response, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from contextlib import asynccontextmanager
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, EmailStr

from services.notification_service import (
    notify_appeal_decision,
    notify_case_closed,
    notify_hearing_scheduled,
    notify_higher_court_reopen_broadcast,
    notify_verdict_issued,
    schedule_case_notifications,
)

# ---------------- Setup ----------------
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
GRAPHS_DIR = ROOT_DIR / "graphs"
GRAPHS_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR = ROOT_DIR / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)

mongo_url = os.getenv("MONGO_URI")
client = AsyncIOMotorClient(mongo_url)
db = client[os.getenv("DB_NAME")]

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# IPFS Configuration
IPFS_DAEMON_ADDRESS = os.getenv("IPFS_DAEMON_ADDRESS", "/ip4/127.0.0.1/tcp/5001/http")
IPFS_GATEWAY_URL = os.getenv("IPFS_GATEWAY_URL", "http://127.0.0.1:8080/ipfs")
IPFS_HTTP_API = os.getenv("IPFS_HTTP_API", "").strip()  # if empty, derived from multiaddr
IPFS_CONNECT_TIMEOUT_S = float(os.getenv("IPFS_CONNECT_TIMEOUT_S", "2.5"))
IPFS_UPLOAD_TIMEOUT_S = float(os.getenv("IPFS_UPLOAD_TIMEOUT_S", "30"))
IPFS_VERIFY_TIMEOUT_S = float(os.getenv("IPFS_VERIFY_TIMEOUT_S", "10"))

# IPFS Connection State
ipfs_client = None
ipfs_connected = False
ipfs_http_available = False
ipfs_version = None
ipfs_status_msg = "Initializing..."
_ipfs_last_check_mono = 0.0
_ipfs_min_recheck_interval_s = float(os.getenv("IPFS_RECHECK_INTERVAL_S", "5"))

def _multiaddr_to_http_api(maddr: str) -> Optional[str]:
    """
    Convert a minimal IPFS HTTP multiaddr to an HTTP base URL.
    Supports patterns like: /ip4/127.0.0.1/tcp/5001/http
    """
    if not maddr:
        return None
    maddr = maddr.strip()
    parts = [p for p in maddr.split("/") if p]
    # expected: ["ip4","127.0.0.1","tcp","5001","http"]
    try:
        if len(parts) >= 5 and parts[0] == "ip4" and parts[2] == "tcp":
            host = parts[1]
            port = parts[3]
            return f"http://{host}:{port}"
    except Exception:
        return None
    return None

def _ipfs_api_base() -> str:
    derived = _multiaddr_to_http_api(IPFS_DAEMON_ADDRESS)
    if IPFS_HTTP_API:
        return IPFS_HTTP_API
    if derived:
        return derived
    # last resort default
    return "http://127.0.0.1:5001"

def _ipfs_post(path: str, *, params: Optional[dict] = None, files=None, timeout_s: float = 10):
    base = _ipfs_api_base().rstrip("/")
    url = f"{base}{path}"
    return requests.post(url, params=params, files=files, timeout=timeout_s)

def _get_ipfs_version() -> Optional[str]:
    """Get IPFS daemon version via HTTP API."""
    try:
        resp = _ipfs_post("/api/v0/version", timeout_s=IPFS_CONNECT_TIMEOUT_S)
        if resp.ok:
            data = resp.json()
            return data.get("Version")
    except Exception:
        pass
    return None

def _ipfs_http_available() -> bool:
    """Check if IPFS HTTP API is available."""
    try:
        resp = _ipfs_post("/api/v0/version", timeout_s=IPFS_CONNECT_TIMEOUT_S)
        return resp.ok
    except Exception:
        return False

def init_ipfs_client():
    """Initialize and test IPFS daemon connection (HTTP API-first for compatibility)."""
    global ipfs_client, ipfs_connected, ipfs_http_available, ipfs_version, ipfs_status_msg, _ipfs_last_check_mono
    
    # First check if HTTP API is available
    ipfs_http_available = _ipfs_http_available()
    if not ipfs_http_available:
        ipfs_status_msg = "IPFS Offline"
        logger.warning(f"✗ IPFS daemon not responding. Evidence will be stored locally only.")
        ipfs_connected = False
        ipfs_client = None
        ipfs_version = None
        _ipfs_last_check_mono = time.monotonic()
        return False
    
    # Get daemon version via HTTP API
    ipfs_version = _get_ipfs_version()
    if ipfs_version:
        logger.info(f"✓ IPFS daemon version: {ipfs_version}")
    
    # For IPFS 0.14.0 compatibility, prefer the HTTP API even if ipfshttpclient exists.
    ipfs_client = None
    ipfs_connected = True
    ipfs_status_msg = f"Connected" + (f" (v{ipfs_version})" if ipfs_version else "")
    logger.info(f"✓ Using IPFS HTTP API for operations at {_ipfs_api_base()}")
    _ipfs_last_check_mono = time.monotonic()
    return True


def refresh_ipfs_state(force: bool = False) -> bool:
    """
    Refresh IPFS connectivity state (cheap, throttled).
    Keeps uploads resilient when daemon toggles online/offline after startup.
    """
    global _ipfs_last_check_mono
    now_mono = time.monotonic()
    if not force and (now_mono - _ipfs_last_check_mono) < _ipfs_min_recheck_interval_s:
        return ipfs_connected and ipfs_http_available
    try:
        return init_ipfs_client()
    except Exception:
        # init_ipfs_client already sets status fields on normal failures,
        # but guard against unexpected exceptions to keep uploads non-fatal.
        _ipfs_last_check_mono = now_mono
        return False

# Initialize IPFS on startup
init_ipfs_client()

# App and API router will be defined after lifespan function
# app = FastAPI(title="AI Criminal Case Management System", lifespan=lifespan)
api = APIRouter(prefix="/api")

JWT_SECRET = os.getenv("JWT_SECRET")
JWT_ALG = "HS256"
ROLES = {"citizen", "police", "forensic", "investigator", "court_officer", "judge", "admin"}

# ---------------- Helpers ----------------
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def hash_password(pwd: str) -> str:
    return bcrypt.hashpw(pwd.encode(), bcrypt.gensalt()).decode()

def verify_password(pwd: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pwd.encode(), hashed.encode())
    except Exception:
        return False

def create_access_token(uid: str, email: str, role: str) -> str:
    payload = {
        "sub": uid, "email": email, "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=12),
        "type": "access",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

def create_refresh_token(uid: str) -> str:
    payload = {
        "sub": uid,
        "exp": datetime.now(timezone.utc) + timedelta(days=7),
        "type": "refresh",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

def set_auth_cookies(response: Response, access: str, refresh: str):
    response.set_cookie("access_token", access, httponly=True, secure=True, samesite="none", max_age=43200, path="/")
    response.set_cookie("refresh_token", refresh, httponly=True, secure=True, samesite="none", max_age=604800, path="/")

def sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _forensic_template_path() -> Path:
    return TEMPLATE_DIR / "sample_forensic_report_template.docx"


def ensure_forensic_report_template() -> Path:
    path = _forensic_template_path()
    if path.exists():
        return path
    if DocxDocument is None:
        fallback = TEMPLATE_DIR / "sample_forensic_report_template.txt"
        fallback.write_text(
            "FORENSIC REPORT TEMPLATE\n\n"
            "Case ID:\nFIR ID:\nExaminer Name:\nLab Name:\nEvidence IDs:\n\n"
            "Hash Verification Section:\n\nCID Verification Section:\n\n"
            "Chain of Custody Section:\n\nAI Tampering Analysis:\n\nFindings:\n\n"
            "Conclusion:\n\nDigital Signature:\nDate & Timestamp:\n",
            encoding="utf-8",
        )
        return fallback
    doc = DocxDocument()
    doc.add_heading("Sample Forensic Report Template", level=1)
    doc.add_paragraph("Case ID: ____________________")
    doc.add_paragraph("FIR ID: _____________________")
    doc.add_paragraph("Examiner Name: _____________")
    doc.add_paragraph("Lab Name: __________________")
    doc.add_paragraph("Evidence IDs: ______________")
    doc.add_heading("Hash Verification Section", level=2)
    doc.add_paragraph("List each evidence hash and verification outcome.")
    doc.add_heading("CID Verification Section", level=2)
    doc.add_paragraph("List each CID and retrieval verification outcome.")
    doc.add_heading("Chain of Custody Section", level=2)
    doc.add_paragraph("Document custody timeline and responsible officers.")
    doc.add_heading("AI Tampering Analysis", level=2)
    doc.add_paragraph("Summarize AI-assisted tampering risk findings.")
    doc.add_heading("Findings", level=2)
    doc.add_paragraph("Summarize forensic findings.")
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("Final conclusion and recommendation.")
    doc.add_paragraph("Digital Signature: ____________________")
    doc.add_paragraph("Date & Timestamp: _____________________")
    doc.save(str(path))
    return path


def _evidence_chain_hash(previous_hash: str, evidence_hash: str, timestamp: str) -> str:
    return sha256_hex(f"{previous_hash}{evidence_hash}{timestamp}".encode())


def _evidence_chain_valid(ev: dict) -> bool:
    if not ev or not ev.get("previous_hash") or not ev.get("timestamp") or not ev.get("current_hash"):
        return False
    return _evidence_chain_hash(ev["previous_hash"], ev["sha256_hash"], ev["timestamp"]) == ev["current_hash"]


def generate_wallet_id() -> str:
    seed = f"{uuid.uuid4()}|{now_iso()}|{os.urandom(16).hex()}"
    return sha256_hex(seed.encode())


def _upload_to_ipfs_http(file_name: str, content: bytes, only_hash: bool = False) -> Optional[str]:
    api_url = f"{_ipfs_api_base().rstrip('/')}/api/v0/add"
    params = {"only-hash": "true"} if only_hash else {"pin": "true"}
    files = {"file": (file_name, content)}

    try:
        resp = requests.post(api_url, params=params, files=files, timeout=IPFS_UPLOAD_TIMEOUT_S)
        if not resp.ok:
            logger.error(f"IPFS HTTP upload failed for {file_name}: {resp.status_code} {resp.text}")
            return None

        for line in resp.text.splitlines():
            if not line.strip():
                continue
            try:
                data = json.loads(line)
            except json.JSONDecodeError:
                continue
            cid = data.get("Hash") or data.get("hash")
            if cid:
                logger.info(f"Generated CID for {file_name} via IPFS HTTP API: {cid}")
                return cid

        logger.error(f"IPFS HTTP upload returned no CID for {file_name}")
        return None
    except requests.Timeout:
        logger.warning(f"IPFS upload timed out for {file_name} after {IPFS_UPLOAD_TIMEOUT_S}s")
        return None
    except Exception as exc:
        logger.error(f"IPFS HTTP upload failed for {file_name}: {exc}")
        return None


def upload_to_ipfs(file_name: str, content: bytes, only_hash: bool = False) -> Optional[str]:
    """
    Upload file content to IPFS daemon.
    Attempts ipfshttpclient first, falls back to HTTP API if needed.
    
    Args:
        file_name: Name of the file being uploaded
        content: File content as bytes
        only_hash: If True, only compute hash without uploading to IPFS
        
    Returns:
        CID (Content Identifier) of the uploaded file, or None if upload fails
    """
    # Ensure daemon status is up-to-date (daemon may start/stop after backend boot).
    refresh_ipfs_state()

    if not ipfs_connected:
        logger.debug(f"IPFS not available. Cannot upload {file_name}")
        return None

    try:
        logger.info(f"{'Generating CID' if only_hash else 'Uploading'} {file_name} via IPFS HTTP API ({len(content)} bytes)")
        if not ipfs_http_available:
            return None
        cid = _upload_to_ipfs_http(file_name, content, only_hash=only_hash)
        if not cid:
            return None
        return cid

    except Exception as exc:
        logger.error(f"IPFS upload error for {file_name}: {exc}")
        return None


def generate_cid_from_content(file_name: str, content: bytes) -> Optional[str]:
    """
    Generate CID for file content without uploading to IPFS.
    Useful for verification purposes.
    
    Returns:
        CID of the content, or None if generation fails
    """
    return upload_to_ipfs(file_name, content, only_hash=True)


def verify_cid_on_ipfs(cid: str) -> bool:
    """
    Verify that a CID exists on the IPFS network.
    
    Args:
        cid: The CID to verify
        
    Returns:
        True if CID exists, False otherwise
    """
    if not ipfs_connected or not ipfs_http_available:
        logger.warning(f"IPFS daemon not connected. Cannot verify CID: {cid}")
        return False

    try:
        resp = _ipfs_post("/api/v0/object/stat", params={"arg": cid}, timeout_s=IPFS_VERIFY_TIMEOUT_S)
        if resp.ok:
            logger.info(f"✓ CID verified on IPFS via HTTP API: {cid}")
            return True
        logger.warning(f"CID HTTP verification failed for {cid}: {resp.status_code} {resp.text}")
        return False
    except requests.Timeout:
        logger.warning(f"CID verification timed out for {cid} after {IPFS_VERIFY_TIMEOUT_S}s")
        return False
    except Exception as exc:
        logger.warning(f"CID verification failed for {cid}: {exc}")
        return False

async def _find_fir_by_relation(fir_relation: str) -> Optional[dict]:
    if not fir_relation:
        return None
    fir_id = fir_relation.strip().split()[0]
    if not fir_id:
        return None
    fir = await db.firs.find_one({"fir_id": fir_id}, {"_id": 0})
    if fir:
        return fir
    return await db.firs.find_one({"fir_id": fir_relation.strip()}, {"_id": 0})

async def _create_case_for_fir(fir_id: str, user: dict) -> dict:
    fir = await db.firs.find_one({"fir_id": fir_id}, {"_id": 0})
    if not fir:
        raise HTTPException(404, "FIR not found")
    existing_case = await db.cases.find_one({"fir_id": fir_id}, {"_id": 0})
    if existing_case:
        return existing_case

    case_id = f"CASE-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"
    start_status = fir.get("status") if fir.get("status") in CASE_STATUSES else "FIR Registered"
    case_doc = {
        "id": str(uuid.uuid4()),
        "case_id": case_id,
        "fir_id": fir_id,
        "citizen_id": fir.get("citizen_id"),
        "citizen_name": fir.get("citizen_name"),
        "title": fir.get("crime_type", "General") + " Case",
        "summary": fir.get("description", ""),
        "crime_type": fir.get("crime_type", ""),
        "location": fir.get("location", ""),
        "status": start_status,
        "status_history": [{"status": start_status, "at": now_iso(), "by": user["email"], "note": "Case created"}],
        "investigation_updates": [],
        "linked_suspects": [],
        "accused_suspects": [],
        "created_by": user["email"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.cases.insert_one(case_doc)
    await append_block(case_id, "CASE_CREATED", {"fir_id": fir_id, "status": start_status}, user["email"])
    await log_activity(user["email"], "case_create", case_id)
    case_doc.pop("_id", None)
    return case_doc

async def _ensure_case_for_suspect_by_fir(suspect_id: str, fir_relation: str, user: dict) -> Optional[str]:
    fir = await _find_fir_by_relation(fir_relation)
    if not fir:
        return None
    case_doc = await _create_case_for_fir(fir["fir_id"], user)
    case_id = case_doc["case_id"]
    now = now_iso()
    await db.suspects.update_one({"suspect_id": suspect_id}, {"$addToSet": {"associated_cases": case_id}})
    await db.cases.update_one(
        {"case_id": case_id},
        {
            "$set": {"updated_at": now},
            "$addToSet": {"linked_suspects": suspect_id},
            "$push": {
                "investigation_updates": {
                    "type": "suspect_auto_linked",
                    "suspect_id": suspect_id,
                    "at": now,
                    "by": user["email"],
                    "note": "Automatically linked from FIR relation",
                }
            },
        },
    )
    await append_block(case_id, "SUSPECT_AUTO_LINKED", {"suspect_id": suspect_id, "fir_id": fir["fir_id"]}, user["email"])
    await log_activity(user["email"], "suspect_auto_link", f"{suspect_id}::{case_id}")
    return case_id

# ---------------- Auth Dependency ----------------
async def get_current_user(request: Request) -> dict:
    token = request.cookies.get("access_token")
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        raise HTTPException(401, "Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        if payload.get("type") != "access":
            raise HTTPException(401, "Invalid token type")
        user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password_hash": 0})
        if not user:
            raise HTTPException(401, "User not found")
        if (user.get("status") or "active").lower() == "suspended":
            raise HTTPException(403, "Account suspended")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Invalid token")

def require_roles(*roles: str):
    """Strict role check: callers must list every role allowed (no implicit admin bypass)."""

    async def checker(user: dict = Depends(get_current_user)):
        if not roles:
            raise HTTPException(500, "require_roles invoked without roles")
        if user["role"] not in roles:
            raise HTTPException(403, f"Requires one of roles: {roles}")
        return user

    return checker

# ---------------- Activity Log ----------------
async def _get_or_create_wallet_by_email(user_email: str) -> Optional[str]:
    if not user_email:
        return None
    user_doc = await db.users.find_one({"email": user_email}, {"_id": 0, "id": 1, "wallet_id": 1})
    if not user_doc:
        return None
    wallet_id = user_doc.get("wallet_id")
    if wallet_id:
        return wallet_id
    wallet_id = generate_wallet_id()
    await db.users.update_one({"id": user_doc["id"]}, {"$set": {"wallet_id": wallet_id}})
    return wallet_id


async def log_activity(user_email: str, action: str, details: str = ""):
    wallet_id = await _get_or_create_wallet_by_email(user_email)
    await db.activity_logs.insert_one({
        "id": str(uuid.uuid4()),
        "user_email": user_email,
        "wallet_id": wallet_id,
        "action": action,
        "details": details,
        "timestamp": now_iso(),
    })


async def flag_evidence_tampering(evidence_id: str, case_id: str, detected_by: str, reason: str):
    await db.fraud_flags.insert_one({
        "id": str(uuid.uuid4()),
        "category": "evidence_tampering",
        "evidence_id": evidence_id,
        "case_id": case_id,
        "detected_by": detected_by,
        "reason": reason,
        "created_at": now_iso(),
    })

# ---------------- Blockchain ----------------
def _normalize_block_data(data: dict) -> dict:
    """Create a detached canonical copy to prevent post-write mutation issues."""
    return json.loads(json.dumps(deepcopy(data or {}), sort_keys=True, separators=(",", ":"), default=str))


def _build_block_hash_payload(case_id: str, action: str, data: dict, timestamp: str, previous_hash: str) -> bytes:
    payload = {
        "caseId": case_id,
        "action": action,
        "data": data,
        "timestamp": timestamp,
        "previousHash": previous_hash,
    }
    return json.dumps(payload, sort_keys=True, separators=(",", ":"), default=str).encode()


def _generate_block_hash(case_id: str, action: str, data: dict, timestamp: str, previous_hash: str) -> str:
    return sha256_hex(_build_block_hash_payload(case_id, action, data, timestamp, previous_hash))


def _get_previous_hash_from_block(block: dict) -> str:
    return block.get("previous_hash") or block.get("previousHash") or ("0" * 64)


def _score_crime_risk(city: str, area: str) -> int:
    city = (city or "").strip().lower()
    area = (area or "").strip().lower()
    score = 0
    keywords = {
        "downtown": 4,
        "market": 3,
        "harbor": 3,
        "port": 3,
        "industrial": 3,
        "old town": 3,
        "station": 2,
        "mall": 2,
        "east": 1,
        "west": 1,
        "south": 1,
        "north": 1,
        "suburb": -1,
        "residential": -1,
        "park": -1,
        "garden": -1,
        "school": -2,
        "university": -2,
        "safe": -3,
    }
    for keyword, weight in keywords.items():
        if keyword in area or keyword in city:
            score += weight
    if any(token in city for token in ["metro", "capital", "city", "downtown"]):
        score += 1
    if not city and not area:
        score = 0
    return score


def _normalize_ai_text(text: str) -> str:
    return (text or "").strip().lower()


def _predict_crime_risk(city: str, area: str) -> dict:
    score = _score_crime_risk(city, area)
    if score >= 4:
        risk_level = "High"
        predicted_crimes = ["Armed robbery", "Assault", "Vehicle theft"]
    elif score >= 1:
        risk_level = "Medium"
        predicted_crimes = ["Burglary", "Vandalism", "Drug possession"]
    else:
        risk_level = "Low"
        predicted_crimes = ["Pickpocketing", "Minor theft", "Noise complaints"]

    area_label = area or "Unknown area"
    city_label = city or "Unknown city"
    reason = (
        f"Rule-based assessment for {area_label} in {city_label}. "
        f"Area keywords and city profile generated a score of {score}."
    )
    return {
        "city": city_label,
        "area": area_label,
        "risk_level": risk_level,
        "predicted_crimes": predicted_crimes,
        "reason": reason,
    }


def _get_high_risk_areas(city: str) -> list:
    lookup = {
        "metro city": ["Downtown", "Harbor District", "Old Market", "Industrial Quarter"],
        "capital city": ["Central Market", "Government Row", "Railway Station", "Riverfront"],
        "coastal city": ["Port Area", "Boardwalk", "Fishing Docks", "East Harbor"],
    }
    city_key = (city or "").strip().lower()
    return lookup.get(city_key, ["Downtown", "Central Market", "Port District", "Old Town", "West End"])


def _answer_ai_question(question: str, city: str, area: str) -> dict:
    prompt = _normalize_ai_text(question)
    if "risky" in prompt or "which area" in prompt or "danger" in prompt:
        areas = _get_high_risk_areas(city)
        return {
            "title": "Risk Area Insight",
            "answer": f"Based on the current location data, the most likely high-risk areas are: {', '.join(areas)}.",
            "notes": "Areas are derived from simple rule-based heuristics, not real-time crime feeds.",
        }
    if "future" in prompt or "prediction" in prompt or "crime" in prompt:
        result = _predict_crime_risk(city, area)
        return {
            "title": "Future Crime Prediction",
            "answer": f"The model estimates {result['risk_level']} risk for {result['area']} in {result['city']}. Likely crimes include {', '.join(result['predicted_crimes'])}.",
            "notes": "This is a mock AI prediction using simple rule-based logic.",
        }
    return {
        "title": "AI Assistant",
        "answer": "I can help identify high-risk areas and produce a crime prediction summary when you provide a city and area.",
        "notes": "Try asking: 'Which area is risky?' or 'Future crime prediction?'",
    }


def _get_current_hash_from_block(block: dict) -> str:
    return block.get("current_hash") or block.get("hash") or ""


async def append_block(case_id: str, action: str, data: dict, user_email: str):
    # Ensure genesis block exists if chain is empty
    last = await db.blockchain.find_one({}, sort=[("index", -1)], projection={"_id": 0})
    if not last:
        # Create genesis block
        genesis_timestamp = now_iso()
        genesis_data = {"genesis": True, "message": "Blockchain genesis block"}
        genesis_hash = _generate_block_hash("GENESIS", "GENESIS", genesis_data, genesis_timestamp, "0" * 64)
        genesis_block = {
            "id": str(uuid.uuid4()),
            "index": 0,
            "timestamp": genesis_timestamp,
            "case_id": "GENESIS",
            "action_type": "GENESIS",
            "action": "GENESIS",
            "data": genesis_data,
            "user_email": "system",
            "wallet_id": None,
            "previous_hash": "0" * 64,
            "previousHash": "0" * 64,
            "current_hash": genesis_hash,
            "hash": genesis_hash,
        }
        await db.blockchain.insert_one(genesis_block)
        last = genesis_block

    prev_hash = _get_current_hash_from_block(last)
    index = last["index"] + 1
    timestamp = now_iso()
    action_type = action
    wallet_id = await _get_or_create_wallet_by_email(user_email)
    normalized_data = _normalize_block_data(data)
    current_hash = _generate_block_hash(case_id, action, normalized_data, timestamp, prev_hash)
    block = {
        "id": str(uuid.uuid4()),
        "index": index,
        "timestamp": timestamp,
        "case_id": case_id,
        "action_type": action_type,
        "action": action,
        "data": normalized_data,
        "user_email": user_email,
        "wallet_id": wallet_id,
        "previous_hash": prev_hash,
        "previousHash": prev_hash,
        "current_hash": current_hash,
        "hash": current_hash,
    }
    await db.blockchain.insert_one(block)
    return block

# ---------------- Pydantic Models ----------------
class RegisterIn(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)
    name: str
    cnic: Optional[str] = ""
    phone: Optional[str] = ""
    role: str = "citizen"

class LoginIn(BaseModel):
    email: EmailStr
    password: str

class FIRCreate(BaseModel):
    crime_type: str
    location: str
    description: str
    incident_date: Optional[str] = None

class FIRDocument(BaseModel):
    filename: str
    file_type: str
    sha256_hash: str
    cid: Optional[str] = None
    uploaded_at: str
    uploaded_by: str

class FIRStatusUpdate(BaseModel):
    status: str
    assigned_officer_id: Optional[str] = None
    note: Optional[str] = ""

class CaseCreate(BaseModel):
    fir_id: str
    title: Optional[str] = ""
    summary: Optional[str] = ""

class CaseStatusUpdate(BaseModel):
    status: str
    # Backwards compatible: frontend historically used `note`.
    note: Optional[str] = ""
    # New fields for workflow routing.
    forwarded_to: Optional[str] = None
    remarks: Optional[str] = ""

class CaseAssignInvestigatorIn(BaseModel):
    investigator_id: str
    note: Optional[str] = ""

class SuspectCreate(BaseModel):
    name: str
    father_name: Optional[str] = ""
    age: Optional[int] = None
    cnic: Optional[str] = ""
    address: Optional[str] = ""
    phone: Optional[str] = ""
    crime_relation: Optional[str] = ""
    fir_relation: Optional[str] = ""
    charges_under_ppc: Optional[List[str]] = []
    arrest_status: str = "not_arrested"  # not_arrested, arrested, on_bail, absconding
    warrant_record: Optional[str] = ""
    witness_statements: Optional[str] = ""
    custody_record: Optional[str] = ""
    investigation_notes: Optional[str] = ""
    crime_history: Optional[str] = ""
    associated_cases: Optional[List[str]] = []
    risk_level: str = "low"  # low, medium, high
    notes: Optional[str] = ""

class SuspectUpdate(BaseModel):
    name: Optional[str] = None
    father_name: Optional[str] = None
    age: Optional[int] = None
    cnic: Optional[str] = None
    address: Optional[str] = None
    phone: Optional[str] = None
    crime_relation: Optional[str] = None
    fir_relation: Optional[str] = None
    charges_under_ppc: Optional[List[str]] = None
    arrest_status: Optional[str] = None
    warrant_record: Optional[str] = None
    witness_statements: Optional[str] = None
    custody_record: Optional[str] = None
    investigation_notes: Optional[str] = None
    crime_history: Optional[str] = None
    associated_cases: Optional[List[str]] = None
    risk_level: Optional[str] = None
    notes: Optional[str] = None

class SuspectAccusedIn(BaseModel):
    case_id: Optional[str] = None
    note: Optional[str] = ""

class SuspectDocumentUploadIn(BaseModel):
    description: Optional[str] = ""

class SuspectVerifyIn(BaseModel):
    verdict: str
    note: Optional[str] = ""

class ForensicCreate(BaseModel):
    case_id: str
    call_logs: Optional[str] = ""
    emails: Optional[str] = ""
    ip_addresses: Optional[str] = ""
    device_ids: Optional[str] = ""
    browser_history: Optional[str] = ""
    social_media: Optional[str] = ""
    deleted_files: Optional[str] = ""
    notes: Optional[str] = ""

class ForensicReportUploadIn(BaseModel):
    case_id: str
    report_title: str
    summary: str
    examiner_name: Optional[str] = ""
    lab_name: Optional[str] = ""
    verified_evidence_ids: Optional[List[str]] = []
    matching_hashes: Optional[List[str]] = []
    ai_tampering_analysis: Optional[str] = ""
    final_conclusion: Optional[str] = ""
    digital_signature: Optional[str] = ""
    image_analysis: Optional[str] = ""
    log_analysis: Optional[str] = ""
    ip_analysis: Optional[str] = ""
    call_data_analysis: Optional[str] = ""
    result: str = "Verified"
    send_to_court: bool = True

class ForensicVerifyEvidenceIn(BaseModel):
    evidence_id: str
    verdict: str = "Verified"  # Verified or Suspicious
    note: Optional[str] = ""

class CourtScheduleHearingIn(BaseModel):
    case_id: str
    hearing_date: str
    note: Optional[str] = ""

class CourtStatusUpdateIn(BaseModel):
    case_id: str
    hearing_status: str  # Scheduled, In Progress, Completed
    note: Optional[str] = ""

class JudgeCloseCaseIn(BaseModel):
    case_id: str
    note: Optional[str] = ""

class JudgeEvidenceActionIn(BaseModel):
    evidence_id: str
    action: str  # verify_hash, verify_cid, approve, reject
    note: Optional[str] = ""

class JudgeCaseReturnIn(BaseModel):
    case_id: str
    target: str  # investigator, forensic
    note: Optional[str] = ""

class AppealDecisionIn(BaseModel):
    decision: str  # accept, reject
    note: Optional[str] = ""

class UserUpdate(BaseModel):
    trust_score: Optional[float] = None
    role: Optional[str] = None


class UserNotificationSettingsIn(BaseModel):
    phone: Optional[str] = None
    notify_sms_enabled: Optional[bool] = None


class AdminUserRoleUpdateIn(BaseModel):
    user_id: str
    role: str


class AdminUserStatusUpdateIn(BaseModel):
    user_id: str
    status: str


class AdminQueueRoleIn(BaseModel):
    user_id: str
    pending_role: str


class AdminTargetUserIn(BaseModel):
    user_id: str


class AdminCaseAssignIn(BaseModel):
    case_id: str
    investigator_id: str


class WitnessCreateIn(BaseModel):
    case_id: str
    fir_id: Optional[str] = None
    name: str
    contact_info: Optional[str] = None
    statement: str
    is_protected: bool = False
    is_confidential: bool = False


class WitnessUpdateIn(BaseModel):
    name: Optional[str] = None
    contact_info: Optional[str] = None
    statement: Optional[str] = None
    is_protected: Optional[bool] = None
    is_confidential: Optional[bool] = None


class WitnessVerifyIn(BaseModel):
    witness_id: str
    hash: Optional[str] = None
    cid: Optional[str] = None


class AdminCaseStatusUpdateIn(BaseModel):
    case_id: str
    status: str


class ConsensusReplaceUserIn(BaseModel):
    case_id: str
    role: str

# ---------------- Auth Routes ----------------
@api.post("/auth/register")
async def register(body: RegisterIn, response: Response):
    email = body.email.lower()
    if body.role not in ROLES:
        raise HTTPException(400, "Invalid role")
    if await db.users.find_one({"email": email}):
        raise HTTPException(400, "Email already registered")
    uid = str(uuid.uuid4())
    user_doc = {
        "id": uid,
        "wallet_id": generate_wallet_id(),
        "email": email,
        "password_hash": hash_password(body.password),
        "name": body.name,
        "cnic": body.cnic,
        "phone": body.phone,
        "role": body.role,
        "trust_score": 100.0,
        "complaints": 0,
        "notify_sms_enabled": True,
        "created_at": now_iso(),
    }
    await db.users.insert_one(user_doc)
    access = create_access_token(uid, email, body.role)
    refresh = create_refresh_token(uid)
    set_auth_cookies(response, access, refresh)
    await log_activity(email, "register", f"role={body.role}")
    user_out = {k: v for k, v in user_doc.items() if k not in ("password_hash", "_id")}
    return user_out

@api.post("/auth/login")
async def login(body: LoginIn, response: Response):
    email = body.email.lower()
    user = await db.users.find_one({"email": email})
    if not user or not verify_password(body.password, user["password_hash"]):
        raise HTTPException(401, "Invalid credentials")
    if (user.get("status") or "active").lower() == "suspended":
        raise HTTPException(403, "Account suspended. Contact an administrator.")
    access = create_access_token(user["id"], email, user["role"])
    refresh = create_refresh_token(user["id"])
    set_auth_cookies(response, access, refresh)
    await log_activity(email, "login")
    user.pop("password_hash", None)
    user.pop("_id", None)
    return user

@api.post("/auth/logout")
async def logout(response: Response, user: dict = Depends(get_current_user)):
    response.delete_cookie("access_token", path="/")
    response.delete_cookie("refresh_token", path="/")
    await log_activity(user["email"], "logout")
    return {"ok": True}

@api.get("/auth/me")
async def me(user: dict = Depends(get_current_user)):
    return user


@api.get("/user/wallet")
async def get_user_wallet(user: dict = Depends(get_current_user)):
    wallet_id = user.get("wallet_id")
    if not wallet_id:
        wallet_id = await _get_or_create_wallet_by_email(user["email"])
    return {"wallet_id": wallet_id}


@api.patch("/user/notification-settings")
async def patch_user_notification_settings(body: UserNotificationSettingsIn, user: dict = Depends(get_current_user)):
    updates = {}
    if body.phone is not None:
        updates["phone"] = (body.phone or "").strip()
    if body.notify_sms_enabled is not None:
        updates["notify_sms_enabled"] = bool(body.notify_sms_enabled)
    if not updates:
        raise HTTPException(400, "No fields to update")
    await db.users.update_one({"id": user["id"]}, {"$set": updates})
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "password_hash": 0})
    return fresh


@api.get("/user/notification-logs")
async def list_user_notification_logs(limit: int = 50, user: dict = Depends(get_current_user)):
    lim = max(1, min(int(limit), 200))
    q = {"$or": [{"recipient_email": user["email"]}, {"recipient_user_id": user["id"]}]}
    rows = await db.notification_logs.find(q, {"_id": 0}).sort("created_at", -1).limit(lim).to_list(lim)
    return {"logs": rows}


# ---------------- FIR Routes ----------------
FIR_STATUSES = [
    "FIR Registered", "Approved", "Under Investigation",
    "Evidence Collected", "Forensic Review", "Sent to Court",
    "Hearing Scheduled", "Judgment Issued", "Closed", "Rejected", "REOPENED",
]

CASE_STATUSES = [
    "FIR Registered", "Approved", "Under Investigation",
    "Evidence Collected", "Forensic Review", "Forensic Review Completed", "Sent to Forensic Review", "Sent to Court",
    "Hearing Scheduled", "Judgment Issued", "Closed", "REOPENED",
]

# Accept snake_case workflow statuses without breaking existing Title Case statuses.
# These aliases are used only for input normalization; the stored status remains the existing Title Case values.
CASE_STATUS_ALIASES = {
    "pending": "FIR Registered",
    "under_investigation": "Under Investigation",
    "forwarded_to_forensic": "Sent to Forensic Review",
    "forensic_review": "Forensic Review",
    "forwarded_to_court": "Sent to Court",
    "in_court": "Hearing Scheduled",
    "verdict_pending": "Judgment Issued",
    "closed": "Closed",
}

def normalize_case_status(raw_status: str) -> str:
    s = (raw_status or "").strip()
    if not s:
        return s
    lowered = s.lower()
    if lowered in CASE_STATUS_ALIASES:
        return CASE_STATUS_ALIASES[lowered]
    return s

def is_case_closed_status(status: str) -> bool:
    s = (status or "").strip().lower()
    return s in ("judgment issued", "closed", "rejected")

def get_display_status(status: str) -> str:
    status = normalize_case_status(status)
    if is_case_closed_status(status):
        return "CASE CLOSED"
    return status or ""

async def sync_fir_status_with_case(case_id: str):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0, "fir_id": 1, "status": 1})
    if not case_doc or not case_doc.get("fir_id"):
        return
    fir_doc = await db.firs.find_one({"fir_id": case_doc["fir_id"]}, {"_id": 0, "status": 1})
    if not fir_doc or fir_doc.get("status") == case_doc.get("status"):
        return
    await db.firs.update_one(
        {"fir_id": case_doc["fir_id"]},
        {
            "$set": {
                "status": case_doc["status"],
                "updated_at": now_iso(),
            },
            "$push": {
                "status_history": {
                    "status": case_doc["status"],
                    "at": now_iso(),
                    "by": "system",
                    "note": "Synchronized from case status",
                }
            },
        },
    )


def validate_case_transition(current_status: str, new_status: str) -> bool:
    current = current_status or ""
    new = new_status or ""
    if current == new:
        return True
    # Investigation → Forensic phase.
    #
    # In practice, some cases start as "Filed" (auto-created) or stay "Under Investigation"
    # while evidence is being uploaded. Frontend already gates "Forward to Forensic" behind
    # evidence presence, so allow forwarding from these early states to prevent 400s.
    if new in ("Forensic Review", "Sent to Forensic Review", "Forensic Review Completed"):
        return current in (
            "Filed",
            "FIR Registered",
            "Approved",
            "Under Investigation",
            "Evidence Collected",
            "Forensic Review",
            "Sent to Forensic Review",
            "Forensic Review Completed",
        )
    # Court phase requires forensic completion.
    if new == "Sent to Court":
        return current in ("Forensic Review Completed", "Sent to Court", "Hearing Scheduled")
    if new == "Hearing Scheduled":
        return current in ("Sent to Court", "Hearing Scheduled")
    # Judgment phase requires court hearing completion.
    if new == "Judgment Issued":
        return current in ("Hearing Scheduled", "Judgment Issued")
    # Case can be closed only after judgment is issued.
    if new == "Closed":
        return current in ("Judgment Issued", "Closed")
    # Higher court appeal accepted — case returns to active lifecycle.
    if current == "REOPENED":
        return new in (
            "REOPENED",
            "Under Investigation",
            "Evidence Collected",
            "Forensic Review",
            "Sent to Forensic Review",
            "Forensic Review Completed",
            "Sent to Court",
            "Hearing Scheduled",
            "Judgment Issued",
            "Closed",
        )
    return True

def ensure_valid_case_transition(current_status: str, new_status: str):
    if not validate_case_transition(current_status, new_status):
        raise HTTPException(400, "Invalid case state transition")

INVESTIGATOR_ALLOWED_STATUS = {
    "Under Investigation",
    "Evidence Collected",
    "Sent to Forensic Review",
}


class InvestigationNoteIn(BaseModel):
    case_id: str
    note: str


class InvestigationLinkSuspectIn(BaseModel):
    case_id: str
    suspect_id: str
    recommended: bool = False
    recommendation_note: Optional[str] = ""


class InvestigationStatusIn(BaseModel):
    case_id: str
    status: str
    update_note: Optional[str] = ""


async def _get_assigned_investigator_case(case_id: str, user: dict) -> dict:
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    assigned_id = case_doc.get("assigned_investigator_id")
    assigned_email = case_doc.get("assigned_investigator_email")
    if assigned_id and assigned_id != user["id"]:
        raise HTTPException(403, "Case not assigned to this investigator")
    if assigned_email and assigned_email != user["email"]:
        raise HTTPException(403, "Case not assigned to this investigator")
    return case_doc

@api.post("/firs")
async def create_fir(
    crime_type: str = Form(...),
    location: str = Form(...),
    description: str = Form(...),
    incident_date: Optional[str] = Form(None),
    files: List[UploadFile] = File(None),
    user: dict = Depends(get_current_user)
):
    fir_id = f"FIR-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"
    documents = []

    # Process uploaded files
    if files:
        for file in files:
            if file.filename:
                # Read file content
                content = await file.read()
                # Calculate SHA-256 hash
                sha256_hash = hashlib.sha256(content).hexdigest()

                # Upload to IPFS if available
                cid = None
                refresh_ipfs_state()
                if ipfs_connected and ipfs_http_available:
                    try:
                        cid = upload_to_ipfs(file.filename or "evidence.bin", content)
                    except Exception as e:
                        logger.warning(f"IPFS upload failed for {file.filename}: {e}")

                # Save file locally as backup
                file_path = UPLOAD_DIR / f"{fir_id}_{file.filename}"
                with open(file_path, "wb") as f:
                    f.write(content)

                documents.append({
                    "filename": file.filename,
                    "file_type": file.content_type or "application/octet-stream",
                    "sha256_hash": sha256_hash,
                    "cid": cid,
                    "uploaded_at": now_iso(),
                    "uploaded_by": user["email"],
                    "file_path": str(file_path)
                })

    doc = {
        "id": str(uuid.uuid4()),
        "fir_id": fir_id,
        "citizen_id": user["id"],
        "citizen_name": user["name"],
        "cnic": user.get("cnic", ""),
        "crime_type": crime_type,
        "location": location,
        "description": description,
        "incident_date": incident_date or now_iso(),
        "status": "FIR Registered",
        "assigned_officer_id": None,
        "assigned_officer_name": None,
        "documents": documents,
        "status_history": [{"status": "FIR Registered", "at": now_iso(), "by": user["email"], "note": ""}],
        "created_at": now_iso(),
    }
    await db.firs.insert_one(doc)

    # Auto-create case for the FIR
    case_id = f"CASE-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"
    case_doc = {
        "id": str(uuid.uuid4()),
        "case_id": case_id,
        "fir_id": fir_id,
        "citizen_id": user["id"],
        "citizen_name": user["name"],
        "title": f"{crime_type} Case",
        "summary": description,
        "crime_type": crime_type,
        "location": location,
        "status": "Filed",
        "status_history": [{"status": "Filed", "at": now_iso(), "by": user["email"], "note": "Case auto-created from FIR"}],
        "created_by": user["email"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.cases.insert_one(case_doc)

    await append_block(case_id, "FIR_CREATED", {"crime_type": crime_type, "location": location, "documents_count": len(documents)}, user["email"])
    await append_block(case_id, "CASE_CREATED", {"fir_id": fir_id, "status": "Filed"}, user["email"])
    await log_activity(user["email"], "fir_create", fir_id)
    await log_activity(user["email"], "case_auto_create", case_id)
    doc.pop("_id", None)
    return doc

@api.get("/firs")
async def list_firs(user: dict = Depends(get_current_user)):
    query = {}
    if user["role"] == "citizen":
        query = {"citizen_id": user["id"]}
    docs = await db.firs.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    for doc in docs:
        doc["display_status"] = get_display_status(doc.get("status"))
    return docs

@api.get("/firs/{fir_id}")
async def get_fir(fir_id: str, user: dict = Depends(get_current_user)):
    fir = await db.firs.find_one({"fir_id": fir_id}, {"_id": 0})
    if not fir:
        raise HTTPException(404, "FIR not found")
    fir["display_status"] = get_display_status(fir.get("status"))
    if user["role"] == "citizen" and fir["citizen_id"] != user["id"]:
        raise HTTPException(403, "Not allowed")
    return fir

@api.patch("/firs/{fir_id}/status")
async def update_fir_status(fir_id: str, body: FIRStatusUpdate, user: dict = Depends(require_roles("police", "forensic"))):
    if body.status not in FIR_STATUSES:
        raise HTTPException(400, "Invalid status")
    fir = await db.firs.find_one({"fir_id": fir_id})
    if not fir:
        raise HTTPException(404, "FIR not found")
    updates = {
        "status": body.status,
        "status_history": fir.get("status_history", []) + [
            {"status": body.status, "at": now_iso(), "by": user["email"], "note": body.note or ""}
        ],
    }
    if body.assigned_officer_id:
        officer = await db.users.find_one({"id": body.assigned_officer_id})
        updates["assigned_officer_id"] = body.assigned_officer_id
        updates["assigned_officer_name"] = officer["name"] if officer else None
    await db.firs.update_one({"fir_id": fir_id}, {"$set": updates})
    await append_block(fir_id, "FIR_STATUS_CHANGE", {"status": body.status}, user["email"])
    await log_activity(user["email"], "fir_status_update", f"{fir_id} → {body.status}")
    updated = await db.firs.find_one({"fir_id": fir_id}, {"_id": 0})
    return updated

# ---------------- Cases Tracking ----------------
@api.post("/cases")
async def create_case(body: CaseCreate, user: dict = Depends(require_roles("police", "forensic"))):
    fir = await db.firs.find_one({"fir_id": body.fir_id}, {"_id": 0})
    if not fir:
        raise HTTPException(404, "FIR not found")
    existing = await db.cases.find_one({"fir_id": body.fir_id}, {"_id": 0})
    if existing:
        raise HTTPException(400, "Case already exists for this FIR")
    case_id = f"CASE-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"
    start_status = fir.get("status") if fir.get("status") in CASE_STATUSES else "FIR Registered"
    doc = {
        "id": str(uuid.uuid4()),
        "case_id": case_id,
        "fir_id": body.fir_id,
        "citizen_id": fir.get("citizen_id"),
        "citizen_name": fir.get("citizen_name"),
        "title": body.title or f"{fir.get('crime_type', 'General')} Case",
        "summary": body.summary or fir.get("description", ""),
        "crime_type": fir.get("crime_type", ""),
        "location": fir.get("location", ""),
        "status": start_status,
        "status_history": [{"status": start_status, "at": now_iso(), "by": user["email"], "note": "Case created"}],
        "created_by": user["email"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.cases.insert_one(doc)
    await append_block(case_id, "CASE_CREATED", {"fir_id": body.fir_id, "status": start_status}, user["email"])
    await log_activity(user["email"], "case_create", case_id)
    doc.pop("_id", None)
    return doc

@api.get("/cases")
async def list_cases(user: dict = Depends(get_current_user)):
    query = {}
    if user["role"] == "citizen":
        query["citizen_id"] = user["id"]
    cases = await db.cases.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    enriched = []
    for case_doc in cases:
        case_doc["display_status"] = get_display_status(case_doc.get("status"))
        forensic_reports = await db.forensic_reports.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("created_at", -1).to_list(100)
        judgments = await db.judgments.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("issued_at", -1).to_list(50)
        final_decision = case_doc.get("judgment_verdict")
        if not final_decision and judgments:
            final_decision = judgments[0].get("verdict")
        timeline = []
        timeline.extend(case_doc.get("status_history", []))
        timeline.extend(case_doc.get("investigation_updates", []))
        timeline.extend(case_doc.get("forensic_updates", []))
        timeline.extend(case_doc.get("court_updates", []))
        timeline.extend(case_doc.get("judge_updates", []))
        timeline.extend(case_doc.get("appeal_timeline", []))
        timeline.sort(key=lambda x: x.get("at", ""), reverse=True)
        case_view = {
            **case_doc,
            "forensic_reports": forensic_reports,
            "judgments": judgments,
            "court_info": {
                "hearing_date": case_doc.get("hearing_date"),
                "hearing_status": case_doc.get("hearing_status"),
            },
            "judge_info": {
                "verdict": case_doc.get("judgment_verdict") or (judgments[0].get("verdict") if judgments else None),
                "decision_note": case_doc.get("judgment_note") or (judgments[0].get("decision_note") if judgments else None),
            },
            "final_case_decision": final_decision,
            "full_timeline": timeline,
        }
        enriched.append(case_view)
    return enriched

@api.get("/cases/{case_id}")
async def get_case(case_id: str, user: dict = Depends(get_current_user)):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    if user["role"] == "citizen" and case_doc.get("citizen_id") != user["id"]:
        raise HTTPException(403, "Not allowed")
    forensic_reports = await db.forensic_reports.find(
        {"case_id": case_doc["case_id"]},
        {"_id": 0},
    ).sort("created_at", -1).to_list(100)
    judgments = await db.judgments.find(
        {"case_id": case_doc["case_id"]},
        {"_id": 0},
    ).sort("issued_at", -1).to_list(50)
    final_decision = case_doc.get("judgment_verdict")
    if not final_decision and judgments:
        final_decision = judgments[0].get("verdict")
    timeline = []
    timeline.extend(case_doc.get("status_history", []))
    timeline.extend(case_doc.get("investigation_updates", []))
    timeline.extend(case_doc.get("forensic_updates", []))
    timeline.extend(case_doc.get("court_updates", []))
    timeline.extend(case_doc.get("judge_updates", []))
    timeline.extend(case_doc.get("appeal_timeline", []))
    timeline.sort(key=lambda x: x.get("at", ""), reverse=True)
    case_doc["display_status"] = get_display_status(case_doc.get("status"))
    return {
        **case_doc,
        "display_status": case_doc["display_status"],
        "forensic_reports": forensic_reports,
        "judgments": judgments,
        "court_info": {
            "hearing_date": case_doc.get("hearing_date"),
            "hearing_status": case_doc.get("hearing_status"),
        },
        "judge_info": {
            "verdict": case_doc.get("judgment_verdict") or (judgments[0].get("verdict") if judgments else None),
            "decision_note": case_doc.get("judgment_note") or (judgments[0].get("decision_note") if judgments else None),
        },
        "final_case_decision": final_decision,
        "full_timeline": timeline,
    }

@api.patch("/cases/{case_id}/status")
async def update_case_status(
    case_id: str,
    body: CaseStatusUpdate,
    user: dict = Depends(require_roles("police", "forensic", "investigator")),
):
    requested_status = body.status
    status = normalize_case_status(body.status)
    if status not in CASE_STATUSES:
        raise HTTPException(400, f"Invalid status: {requested_status}")
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), status)
    remarks = (body.remarks if body.remarks is not None else body.note) or ""
    remarks = remarks.strip()
    updates = {
        "status": status,
        "updated_at": now_iso(),
        "status_history": case_doc.get("status_history", []) + [
            {"status": status, "at": now_iso(), "by": user["email"], "note": remarks}
        ],
    }
    # Preserve existing workflow flags used elsewhere in dashboards.
    if status == "Sent to Forensic Review":
        updates.update({
            "forwarded_to_forensic": True,
            "forwarded_to_forensic_at": now_iso(),
            "forwarded_to_forensic_by": user["email"],
        })
    if status == "Sent to Court":
        updates.update({
            "forwarded_to_court": True,
            "forwarded_to_court_at": now_iso(),
            "forwarded_to_court_by": user["email"],
        })
    if body.forwarded_to:
        updates["forwarded_to"] = body.forwarded_to
    await db.cases.update_one({"case_id": case_id}, {"$set": updates})
    await append_block(case_id, "CASE_STATUS_CHANGE", {"status": status, "requested_status": requested_status, "forwarded_to": body.forwarded_to, "remarks": remarks}, user["email"])
    await log_activity(user["email"], "case_status_update", f"{case_id} → {status}")
    updated = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    await sync_fir_status_with_case(case_id)
    return updated


@api.get("/investigator/list")
async def list_investigators(user: dict = Depends(require_roles("police"))):
    investigators = await db.users.find(
        {"role": "investigator"},
        {"_id": 0, "password_hash": 0}
    ).to_list(500)

    # normalize response for frontend dropdown
    return [
        {
            "id": inv.get("id"),
            "name": inv.get("name", inv.get("email")),
            "email": inv.get("email")
        }
        for inv in investigators
    ]


@api.patch("/cases/{case_id}/assign-investigator")
async def assign_investigator_to_case(
    case_id: str,
    body: CaseAssignInvestigatorIn,
    user: dict = Depends(require_roles("police")),
):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Under Investigation")
    investigator = await db.users.find_one({"id": body.investigator_id}, {"_id": 0, "password_hash": 0})
    if not investigator or investigator.get("role") != "investigator":
        raise HTTPException(404, "Investigator not found")
    note = (body.note or "").strip() or f"Assigned investigator: {investigator.get('name', investigator['email'])}"
    status_entry = {
        "status": "Under Investigation",
        "at": now_iso(),
        "by": user["email"],
        "note": note,
    }
    await db.cases.update_one(
        {"case_id": case_id},
        {
            "$set": {
                "assigned_investigator_id": investigator["id"],
                "assigned_investigator_name": investigator.get("name", ""),
                "assigned_investigator_email": investigator.get("email", ""),
                "status": "Under Investigation",
                "updated_at": now_iso(),
            },
            "$push": {
                "status_history": status_entry,
                "investigation_updates": {
                    "type": "investigator_assignment",
                    "investigator_id": investigator["id"],
                    "investigator_email": investigator.get("email", ""),
                    "at": now_iso(),
                    "by": user["email"],
                    "note": note,
                },
            },
        },
    )
    await append_block(
        case_id,
        "INVESTIGATOR_ASSIGNED",
        {"investigator_id": investigator["id"], "status": "Under Investigation"},
        user["email"],
    )
    await log_activity(user["email"], "case_assign_investigator", f"{case_id} → {investigator['email']}")
    updated = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    return updated


@api.patch("/firs/{fir_id}/assign-investigator")
async def assign_investigator_by_fir(
    fir_id: str,
    body: CaseAssignInvestigatorIn,
    user: dict = Depends(require_roles("police")),
):
    case_doc = await db.cases.find_one({"fir_id": fir_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found for this FIR")
    return await assign_investigator_to_case(case_doc["case_id"], body, user)


# ---------------- Investigator Module ----------------
@api.get("/investigator/cases")
async def investigator_cases(user: dict = Depends(require_roles("investigator"))):
    query = {
        "$or": [
            {"assigned_investigator_id": user["id"]},
            {"assigned_investigator_email": user["email"]},
        ]
    }
    assigned_cases = await db.cases.find(query, {"_id": 0}).sort("updated_at", -1).to_list(500)
    results = []
    for case_doc in assigned_cases:
        fir = await db.firs.find_one({"fir_id": case_doc["fir_id"]}, {"_id": 0})
        suspects = await db.suspects.find(
            {"associated_cases": {"$in": [case_doc["case_id"]]}},
            {"_id": 0},
        ).to_list(100)
        notes = await db.investigation_notes.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("created_at", -1).to_list(200)
        results.append({
            "case": case_doc,
            "fir": fir,
            "suspects": suspects,
            "notes": notes,
        })
    return results


@api.post("/investigator/add-note")
async def investigator_add_note(body: InvestigationNoteIn, user: dict = Depends(require_roles("investigator"))):
    case_doc = await _get_assigned_investigator_case(body.case_id, user)
    note_text = (body.note or "").strip()
    if not note_text:
        raise HTTPException(400, "Note is required")
    note_doc = {
        "id": str(uuid.uuid4()),
        "case_id": body.case_id,
        "note": note_text,
        "created_by": user["email"],
        "created_at": now_iso(),
    }
    await db.investigation_notes.insert_one(note_doc)
    history_entry = {
        "type": "note",
        "note": note_text,
        "at": now_iso(),
        "by": user["email"],
    }
    await db.cases.update_one(
        {"case_id": case_doc["case_id"]},
        {"$set": {"updated_at": now_iso()}, "$push": {"investigation_updates": history_entry}},
    )
    await append_block(case_doc["case_id"], "INVESTIGATION_NOTE", {"note": note_text}, user["email"])
    await log_activity(user["email"], "investigator_add_note", case_doc["case_id"])
    note_doc.pop("_id", None)
    return note_doc


@api.post("/investigator/link-suspect")
async def investigator_link_suspect(body: InvestigationLinkSuspectIn, user: dict = Depends(require_roles("investigator"))):
    case_doc = await _get_assigned_investigator_case(body.case_id, user)
    suspect = await db.suspects.find_one({"suspect_id": body.suspect_id}, {"_id": 0})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    await db.suspects.update_one(
        {"suspect_id": body.suspect_id},
        {"$addToSet": {"associated_cases": body.case_id}},
    )
    link_entry = {
        "suspect_id": body.suspect_id,
        "linked_at": now_iso(),
        "linked_by": user["email"],
        "recommended": body.recommended,
        "recommendation_note": (body.recommendation_note or "").strip(),
    }
    case_update = {
        "$set": {"updated_at": now_iso()},
        "$addToSet": {"linked_suspects": body.suspect_id},
        "$push": {"investigation_updates": {
            "type": "suspect_linked",
            "suspect_id": body.suspect_id,
            "recommended": body.recommended,
            "note": (body.recommendation_note or "").strip(),
            "at": now_iso(),
            "by": user["email"],
        }},
    }
    if body.recommended:
        case_update["$addToSet"]["recommended_suspects"] = body.suspect_id
    await db.cases.update_one({"case_id": case_doc["case_id"]}, case_update)
    await append_block(case_doc["case_id"], "SUSPECT_LINKED", link_entry, user["email"])
    await log_activity(user["email"], "investigator_link_suspect", f"{case_doc['case_id']}::{body.suspect_id}")
    return {"ok": True, "case_id": body.case_id, "suspect_id": body.suspect_id}


@api.post("/investigator/update-status")
async def investigator_update_status(body: InvestigationStatusIn, user: dict = Depends(require_roles("investigator"))):
    if body.status not in INVESTIGATOR_ALLOWED_STATUS:
        raise HTTPException(400, "Invalid investigator status")
    case_doc = await _get_assigned_investigator_case(body.case_id, user)
    mapped_status = "Forensic Review" if body.status == "Sent to Forensic Review" else body.status
    ensure_valid_case_transition(case_doc.get("status"), mapped_status)
    status_entry = {
        "status": mapped_status,
        "at": now_iso(),
        "by": user["email"],
        "note": (body.update_note or "").strip(),
    }
    await db.cases.update_one(
        {"case_id": case_doc["case_id"]},
        {
            "$set": {"status": mapped_status, "updated_at": now_iso()},
            "$push": {
                "status_history": status_entry,
                "investigation_updates": {
                    "type": "status_update",
                    "status": body.status,
                    "at": now_iso(),
                    "by": user["email"],
                    "note": (body.update_note or "").strip(),
                },
            },
        },
    )
    if body.status == "Sent to Forensic Review":
        await db.cases.update_one(
            {"case_id": case_doc["case_id"]},
            {"$set": {
                "forwarded_to_forensic": True,
                "forwarded_to_forensic_at": now_iso(),
                "forwarded_to_forensic_by": user["email"],
            }},
        )
    await append_block(case_doc["case_id"], "INVESTIGATION_STATUS_UPDATE", {"status": body.status}, user["email"])
    await log_activity(user["email"], "investigator_update_status", f"{case_doc['case_id']} → {body.status}")
    updated = await db.cases.find_one({"case_id": case_doc["case_id"]}, {"_id": 0})
    return updated

# ---------------- Suspects ----------------
@api.post("/suspects")
async def create_suspect(body: SuspectCreate, user: dict = Depends(require_roles("police", "forensic", "investigator"))):
    payload = body.model_dump()
    arrest_status = (payload.get("arrest_status") or "not_arrested").strip().lower()
    if arrest_status not in ("not_arrested", "arrested", "on_bail", "absconding"):
        raise HTTPException(400, "Invalid arrest status")
    payload["arrest_status"] = arrest_status
    doc = {
        "id": str(uuid.uuid4()),
        "suspect_id": f"SUS-{str(uuid.uuid4())[:6].upper()}",
        **payload,
        "is_accused": False,
        "accused_history": [],
        "documents": [],
        "verifications": [],
        "created_by": user["email"],
        "created_at": now_iso(),
    }
    await db.suspects.insert_one(doc)
    if payload.get("fir_relation"):
        await _ensure_case_for_suspect_by_fir(doc["suspect_id"], payload["fir_relation"], user)
    await log_activity(user["email"], "suspect_add", doc["suspect_id"])
    result = await db.suspects.find_one({"suspect_id": doc["suspect_id"]}, {"_id": 0})
    return result or doc

@api.get("/suspects")
async def list_suspects(user: dict = Depends(get_current_user)):
    if user["role"] == "citizen":
        raise HTTPException(403, "Not allowed")
    return await db.suspects.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)

@api.patch("/suspects/{suspect_id}")
async def update_suspect(suspect_id: str, body: SuspectUpdate, user: dict = Depends(require_roles("police", "investigator"))):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if "arrest_status" in updates:
        arrest_status = str(updates.get("arrest_status") or "").strip().lower()
        if arrest_status not in ("not_arrested", "arrested", "on_bail", "absconding"):
            raise HTTPException(400, "Invalid arrest status")
        updates["arrest_status"] = arrest_status

    case_id = None
    if "fir_relation" in updates and updates.get("fir_relation"):
        case_id = await _ensure_case_for_suspect_by_fir(suspect_id, updates["fir_relation"], user)

    if not updates and not case_id:
        return suspect

    update_ops = {}
    if updates:
        updates["updated_at"] = now_iso()
        update_ops["$set"] = updates
    if case_id:
        update_ops.setdefault("$addToSet", {})["associated_cases"] = case_id
    if update_ops:
        await db.suspects.update_one({"suspect_id": suspect_id}, update_ops)
    suspect_anchor = (suspect.get("associated_cases") or [None])[0] or suspect_id
    await append_block(
        suspect_anchor,
        "SUSPECT_UPDATED",
        {"suspect_id": suspect_id, "fields": sorted(list(updates.keys())) if updates else []},
        user["email"],
    )
    await log_activity(user["email"], "suspect_update", suspect_id)
    return await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})

@api.post("/suspects/{suspect_id}/mark-accused")
async def mark_suspect_accused(suspect_id: str, body: SuspectAccusedIn, user: dict = Depends(require_roles("investigator", "court_officer", "judge"))):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    case_id = body.case_id or (suspect.get("associated_cases") or [None])[0]
    if not case_id and suspect.get("fir_relation"):
        case_id = await _ensure_case_for_suspect_by_fir(suspect_id, suspect["fir_relation"], user)
    if not case_id:
        raise HTTPException(400, "No case association found for suspect")
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    accused_entry = {
        "case_id": case_id,
        "at": now_iso(),
        "by": user["email"],
        "note": (body.note or "").strip(),
    }
    await db.suspects.update_one(
        {"suspect_id": suspect_id},
        {
            "$set": {"is_accused": True, "updated_at": now_iso()},
            "$addToSet": {"associated_cases": case_id},
            "$push": {"accused_history": accused_entry},
        },
    )
    await db.cases.update_one(
        {"case_id": case_id},
        {
            "$addToSet": {"accused_suspects": suspect_id},
            "$push": {"investigation_updates": {
                "type": "suspect_marked_accused",
                "suspect_id": suspect_id,
                "at": now_iso(),
                "by": user["email"],
                "note": (body.note or "").strip(),
            }},
            "$set": {"updated_at": now_iso()},
        },
    )
    await append_block(case_id, "SUSPECT_MARKED_ACCUSED", {"suspect_id": suspect_id, "note": (body.note or "").strip()}, user["email"])
    await log_activity(user["email"], "suspect_mark_accused", f"{case_id}::{suspect_id}")
    return {"ok": True, "case_id": case_id, "suspect_id": suspect_id}

@api.post("/suspects/{suspect_id}/documents")
async def upload_suspect_document(
    suspect_id: str,
    description: str = Form(""),
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("police", "investigator", "forensic")),
):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file")
    file_hash = sha256_hex(content)
    refresh_ipfs_state()
    cid = None
    ipfs_status = "offline"
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "suspect_doc.bin", content)
        ipfs_status = "online" if cid else "failed"
    upload_id = str(uuid.uuid4())
    ext = Path(file.filename or "").suffix
    stored_name = f"SUS-{upload_id}{ext}"
    timestamp = now_iso()
    (UPLOAD_DIR / stored_name).write_bytes(content)

    doc = {
        "id": upload_id,
        "document_id": f"SDOC-{str(uuid.uuid4())[:6].upper()}",
        "suspect_id": suspect_id,
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "description": (description or "").strip(),
        "sha256_hash": file_hash,
        "cid": cid,
        "ipfs_status": ipfs_status,
        "uploaded_by": user["email"],
        "uploaded_at": timestamp,
        "chain_of_custody": [
            {
                "action": "Suspect Document Uploaded",
                "performedBy": f"{user.get('name', user['email'])} ({user.get('role', '').capitalize()})",
                "timestamp": timestamp,
                "remarks": f"IPFS status: {ipfs_status}",
            }
        ],
        "verification_approvals": [],
    }
    await db.suspects.update_one(
        {"suspect_id": suspect_id},
        {"$push": {"documents": doc}, "$set": {"updated_at": now_iso()}},
    )
    suspect_anchor = (suspect.get("associated_cases") or [None])[0] or suspect_id
    await append_block(
        suspect_anchor,
        "SUSPECT_DOCUMENT_UPLOADED",
        {"suspect_id": suspect_id, "document_id": doc["document_id"], "sha256_hash": file_hash, "cid": cid},
        user["email"],
    )
    await log_activity(user["email"], "suspect_document_upload", f"{suspect_id}::{doc['document_id']}")
    doc.pop("_id", None)
    return doc

@api.post("/suspects/{suspect_id}/documents/{document_id}/verify")
async def verify_suspect_document(
    suspect_id: str,
    document_id: str,
    user: dict = Depends(require_roles("investigator", "forensic", "judge", "court_officer")),
):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    docs = suspect.get("documents", []) or []
    target = next((d for d in docs if d.get("document_id") == document_id), None)
    if not target:
        raise HTTPException(404, "Document not found")
    if not target.get("sha256_hash") or not target.get("cid"):
        raise HTTPException(400, "Document cannot be verified without hash and CID")

    file_path = UPLOAD_DIR / target.get("stored_filename", "")
    if not file_path.exists():
        raise HTTPException(404, "Stored suspect document file not found for verification")

    current_hash = sha256_hex(file_path.read_bytes())
    hash_verified = current_hash == target["sha256_hash"]
    if not hash_verified:
        return {"ok": False, "success": False, "hash_verified": False, "cid_verified": False, "message": "Tampered or Invalid"}

    refresh_ipfs_state()
    if not (ipfs_connected and ipfs_http_available):
        return {"ok": False, "success": False, "hash_verified": True, "cid_verified": False, "message": "IPFS content temporarily unavailable"}

    cid_verified = verify_cid_on_ipfs(target["cid"])
    if not cid_verified:
        return {"ok": False, "success": False, "hash_verified": True, "cid_verified": False, "message": "Tampered or Invalid"}

    verification = {
        "role": user["role"],
        "verified_by": user["email"],
        "verified_at": now_iso(),
        "hash_verified": True,
        "cid_verified": True,
        "status": "approved",
    }
    chain_entry = {
        "action": "Suspect Document Verified",
        "performedBy": f"{user.get('name', user['email'])} ({user.get('role', '').capitalize()})",
        "timestamp": now_iso(),
        "remarks": "Hash + CID verification approved",
    }
    await db.suspects.update_one(
        {"suspect_id": suspect_id, "documents.document_id": document_id},
        {
            "$push": {
                "documents.$.verification_approvals": verification,
                "documents.$.chain_of_custody": chain_entry,
            },
            "$set": {"updated_at": now_iso()},
        },
    )
    await append_block(
        suspect.get("associated_cases", [None])[0] or suspect_id,
        "SUSPECT_DOCUMENT_VERIFIED",
        {"suspect_id": suspect_id, "document_id": document_id, "verified_by": user["email"], "role": user["role"]},
        user["email"],
    )
    await log_activity(user["email"], "suspect_document_verify", f"{suspect_id}::{document_id}")
    return {"ok": True, "success": True, "hash_verified": True, "cid_verified": True, "message": "Document Authenticated"}

@api.post("/suspects/{suspect_id}/verify")
async def verify_suspect(suspect_id: str, body: SuspectVerifyIn, user: dict = Depends(require_roles("investigator", "forensic", "court_officer", "judge"))):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    if body.verdict not in ("Verified", "Suspicious"):
        raise HTTPException(400, "Invalid verdict")
    verification_doc = {
        "id": str(uuid.uuid4()),
        "suspect_id": suspect_id,
        "verdict": body.verdict,
        "note": (body.note or "").strip(),
        "verified_by": user["email"],
        "verified_at": now_iso(),
    }
    await db.suspects.update_one(
        {"suspect_id": suspect_id},
        {"$push": {"verifications": verification_doc}, "$set": {"updated_at": now_iso()}},
    )
    await log_activity(user["email"], "suspect_verified", f"{suspect_id}::{body.verdict}")
    verification_doc.pop("_id", None)
    return verification_doc

@api.patch("/suspects/{suspect_id}/sentence")
async def update_suspect_sentence(suspect_id: str, body: dict, user: dict = Depends(require_roles("judge", "admin"))):
    suspect = await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})
    if not suspect:
        raise HTTPException(404, "Suspect not found")
    
    sentence_status = body.get("sentence_status")
    if sentence_status not in ("Active Sentence", "Completed Sentence", "Released"):
        raise HTTPException(400, "Invalid sentence status")
    
    updates = {
        "sentence_status": sentence_status,
        "updated_at": now_iso(),
    }
    
    # If releasing, clear expected release date
    if sentence_status == "Released":
        updates["expected_release_date"] = None
    
    await db.suspects.update_one({"suspect_id": suspect_id}, {"$set": updates})
    
    await append_block(
        suspect.get("case_id") or suspect_id,
        "SENTENCE_STATUS_UPDATED",
        {"suspect_id": suspect_id, "sentence_status": sentence_status},
        user["email"],
    )
    await log_activity(user["email"], "sentence_status_update", f"{suspect_id}::{sentence_status}")
    
    return await db.suspects.find_one({"suspect_id": suspect_id}, {"_id": 0})

# ---------------- Evidence (with file upload) ----------------
@api.get("/health/ipfs")
async def check_ipfs_health(user: dict = Depends(get_current_user)):
    """
    Check IPFS daemon health and connectivity.
    Available to all authenticated users.
    """
    refresh_ipfs_state()
    status_info = {
        "status": "online" if ipfs_connected else "offline",
        "connected": ipfs_connected,
        "message": ipfs_status_msg,
        "version": ipfs_version,
        "daemon_address": IPFS_DAEMON_ADDRESS,
        "http_api": _ipfs_api_base(),
        "gateway_url": IPFS_GATEWAY_URL,
    }

    if ipfs_connected:
        status_info["client_type"] = "http_api"
        try:
            resp = _ipfs_post("/api/v0/id", timeout_s=IPFS_CONNECT_TIMEOUT_S)
            if resp.ok:
                ipfs_id = resp.json()
                status_info["peer_id"] = ipfs_id.get("ID", "Unknown")
                status_info["agent_version"] = ipfs_id.get("AgentVersion", "Unknown")
            else:
                status_info["peer_id"] = "unavailable"
        except Exception as e:
            logger.warning(f"IPFS ID fetch failed: {e}")
            status_info["peer_id"] = "unavailable"
        status_info["recommendation"] = "IPFS is ready for evidence uploads with CID generation"
    else:
        status_info["client_type"] = "none"
        status_info["peer_id"] = None
        status_info["recommendation"] = "Start IPFS daemon: ipfs daemon"

    return status_info

@api.get("/ipfs/status")
async def get_ipfs_status():
    """
    Public IPFS status endpoint (no authentication required).
    Used by frontend to display IPFS connection badge.
    """
    refresh_ipfs_state()
    return {
        "connected": ipfs_connected,
        "status": "Connected" if ipfs_connected else "Offline",
        "message": ipfs_status_msg,
        "version": ipfs_version,
        "http_api_available": ipfs_http_available,
        "gateway_url": IPFS_GATEWAY_URL,
        "daemon_address": IPFS_DAEMON_ADDRESS,
    }


@api.post("/ipfs/upload")
async def ipfs_upload(
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("police", "forensic")),
):
    """
    Upload file to IPFS only (not to case).
    Returns CID if successful, or error if IPFS daemon is offline.
    """
    content = await file.read()
    
    if not ipfs_connected:
        raise HTTPException(
            503, 
            detail="IPFS daemon is offline. Cannot upload to IPFS. Please ensure IPFS daemon is running at " + IPFS_DAEMON_ADDRESS
        )
    
    cid = upload_to_ipfs(file.filename or "evidence.bin", content)
    if not cid:
        raise HTTPException(502, detail="IPFS upload failed. Please check IPFS daemon status.")
    
    await log_activity(user["email"], "ipfs_upload", cid)
    logger.info(f"IPFS upload completed by {user['email']}: {file.filename} -> {cid}")
    return {
        "cid": cid,
        "file_name": file.filename,
        "uploaded_by": user["email"],
        "gateway_url": f"{IPFS_GATEWAY_URL}/{cid}"
    }


@api.post("/evidence")
async def upload_evidence(
    case_id: str = Form(...),
    evidence_type: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("police", "forensic", "investigator")),
):
    """
    Upload evidence file with SHA256 hash verification.
    Attempts to upload to IPFS if daemon is available, but file is always stored locally.
    Returns evidence document with CID (if IPFS available) and SHA256 hash.
    """
    content = await file.read()
    file_hash = sha256_hex(content)
    
    # Attempt IPFS upload, but don't fail if IPFS is offline
    cid = None
    ipfs_status = "offline"
    refresh_ipfs_state()
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "evidence.bin", content)
        ipfs_status = "online" if cid else "failed"
        logger.info(f"IPFS upload attempt: {ipfs_status} for {file.filename}")
    else:
        logger.warning(f"IPFS daemon offline - storing evidence locally only: {file.filename}")
    
    eid = str(uuid.uuid4())
    ext = Path(file.filename or "").suffix
    stored_name = f"{eid}{ext}"
    timestamp = now_iso()
    
    # Get the last evidence in the chain for this case
    last_evidence = await db.evidence.find_one(
        {"case_id": case_id, "current_hash": {"$exists": True}},
        sort=[("timestamp", -1)],
        projection={"current_hash": 1},
    )
    previous_hash = last_evidence["current_hash"] if last_evidence else "0" * 64
    current_hash = _evidence_chain_hash(previous_hash, file_hash, timestamp)
    
    # Store file locally
    (UPLOAD_DIR / stored_name).write_bytes(content)
    logger.info(f"Evidence stored locally: {stored_name} (SHA256: {file_hash})")
    
    doc = {
        "id": eid,
        "evidence_id": f"EVD-{str(uuid.uuid4())[:6].upper()}",
        "case_id": case_id,
        "type": evidence_type,
        "description": description,
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "size_bytes": len(content),
        "sha256_hash": file_hash,
        "previous_hash": previous_hash,
        "current_hash": current_hash,
        "uploaded_by": user["email"],
        "wallet_id": await _get_or_create_wallet_by_email(user["email"]),
        "uploaded_at": timestamp,
        "file_name": file.filename,
        "file_type": file.content_type or evidence_type,
        "cid": cid,  # Will be None if IPFS offline
        "ipfs_status": ipfs_status,  # Track IPFS upload status
        "cid_verified": False,
        "cid_verified_by": None,
        "cid_verified_at": None,
        "timestamp": timestamp,
        "tampered": False,
        "approvals": [],
        "chainOfCustody": [
            {
                "action": "Uploaded",
                "performedBy": f"{user['name']} ({user['role'].capitalize()})",
                "timestamp": timestamp,
                "remarks": f"IPFS Status: {ipfs_status}",
            }
        ],
    }
    
    await db.evidence.insert_one(doc)
    
    # Log blockchain event
    await append_block(
        case_id,
        "EVIDENCE_UPLOADED",
        {
            "evidence_id": doc["evidence_id"],
            "hash": file_hash,
            "cid": cid,
            "ipfs_status": ipfs_status
        },
        user["email"],
    )
    
    await log_activity(user["email"], "evidence_upload", doc["evidence_id"])
    logger.info(f"Evidence uploaded: {doc['evidence_id']} by {user['email']} (IPFS: {ipfs_status})")
    
    doc.pop("_id", None)
    return doc

@api.post("/evidence/{evidence_id}/chain-doc")
async def upload_evidence_chain_doc(
    evidence_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("police", "investigator", "forensic")),
):
    ev = await db.evidence.find_one({"evidence_id": evidence_id})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file")
    file_hash = sha256_hex(content)
    refresh_ipfs_state()
    cid = None
    ipfs_status = "offline"
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "chain_doc.bin", content)
        ipfs_status = "online" if cid else "failed"
    document_id = str(uuid.uuid4())
    ext = Path(file.filename or "").suffix
    stored_name = f"COC-{document_id}{ext}"
    timestamp = now_iso()
    (UPLOAD_DIR / stored_name).write_bytes(content)
    doc = {
        "id": document_id,
        "document_id": f"COC-{str(uuid.uuid4())[:6].upper()}",
        "evidence_id": evidence_id,
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "sha256_hash": file_hash,
        "cid": cid,
        "ipfs_status": ipfs_status,
        "uploaded_by": user["email"],
        "uploaded_at": timestamp,
    }
    await db.evidence.update_one(
        {"evidence_id": evidence_id},
        {"$push": {"chain_of_custody_docs": doc}, "$set": {"updated_at": now_iso()}},
    )
    await log_activity(user["email"], "evidence_chain_doc_upload", f"{evidence_id}::{doc['document_id']}")
    doc.pop("_id", None)
    return doc

@api.get("/evidence")
async def list_evidence(case_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    query = {}
    if case_id:
        query["case_id"] = case_id
    if user["role"] == "citizen":
        # citizens only see evidence for their own FIRs
        my_firs = await db.firs.find({"citizen_id": user["id"]}, {"fir_id": 1, "_id": 0}).to_list(500)
        my_ids = [f["fir_id"] for f in my_firs]
        query["case_id"] = {"$in": my_ids}
    docs = await db.evidence.find(query, {"_id": 0}).sort("uploaded_at", -1).to_list(500)
    for doc in docs:
        doc["chain_status"] = "Valid Chain" if _evidence_chain_valid(doc) else "Broken Chain"
        doc["consensus_reached"] = len(doc.get("approvals", [])) >= 2
        doc["cid_verified"] = doc.get("cid_verified", False)
        doc["cid_verified_by"] = doc.get("cid_verified_by")
        doc["cid_verified_at"] = doc.get("cid_verified_at")
    return docs

@api.get("/evidence/{evidence_id}/verify")
async def verify_evidence(evidence_id: str, user: dict = Depends(get_current_user)):
    ev = await db.evidence.find_one({"evidence_id": evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    path = UPLOAD_DIR / ev["stored_filename"]
    if not path.exists():
        return {"ok": False, "reason": "file_missing", "original_hash": ev["sha256_hash"]}
    current_hash = sha256_hex(path.read_bytes())
    ok = current_hash == ev["sha256_hash"]
    entry = {
        "action": "Verified",
        "performedBy": f"{user['name']} ({user['role'].capitalize()})",
        "timestamp": now_iso(),
        "remarks": "Integrity verified" if ok else "Hash mismatch detected",
    }
    await db.evidence.update_one({"evidence_id": evidence_id}, {"$push": {"chainOfCustody": entry}})
    if not ok:
        await db.evidence.update_one({"evidence_id": evidence_id}, {"$set": {"tampered": True}})
        await flag_evidence_tampering(evidence_id, ev.get("case_id", ""), user["email"], "Hash mismatch detected")
        await log_activity(user["email"], "evidence_tampered_detected", evidence_id)
    return {
        "ok": ok,
        "evidence_id": evidence_id,
        "original_hash": ev["sha256_hash"],
        "current_hash": current_hash,
        "chain_status": "Valid Chain" if _evidence_chain_valid(ev) else "Broken Chain",
        "consensus_reached": len(ev.get("approvals", [])) >= 2,
        "message": "Evidence Verified Successfully - No Tampering Detected" if ok else "Warning: Evidence May Be Tampered - Hash Mismatch Detected",
    }


@api.post("/evidence/{evidence_id}/cid-verify")
async def verify_evidence_cid(
    evidence_id: str,
    cid: str = Form(...),
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("forensic", "judge", "court_officer")),
):
    """
    Verify CID of uploaded evidence file.
    Only accessible to: Judge, Court Officer, Forensic Expert
    
    This endpoint:
    1. Generates CID from uploaded file
    2. Compares with entered CID
    3. Compares with stored CID in evidence record
    4. Optionally verifies CID exists on IPFS network
    5. Records verification in chain of custody
    """
    ev = await db.evidence.find_one({"evidence_id": evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    
    content = await file.read()
    
    # Generate CID from uploaded file
    generated_cid = generate_cid_from_content(
        file.filename or ev.get("original_filename", "evidence.bin"),
        content
    )
    
    if generated_cid is None:
        # CID generation failed - treat as graceful offline/failure
        msg = "CID generation failed (IPFS Offline or timed out)."
        logger.warning(f"CID generation failed for {evidence_id}: {msg}")
        return {
            "ok": False,
            "entered_cid": cid,
            "generated_cid": None,
            "expected_cid": ev.get("cid"),
            "file_matches": False,
            "record_matches": cid == ev.get("cid"),
            "ipfs_verified": False,
            "ipfs_status": "Offline" if not ipfs_connected else "Failed",
            "message": msg,
            "role": user["role"],
            "verified_by": user["name"],
        }
    
    # Compare CIDs
    file_matches = generated_cid == cid  # Does file match entered CID?
    record_matches = cid == ev.get("cid")  # Does entered CID match stored CID?
    
    # Both must match for verification to succeed
    verification_ok = file_matches and record_matches
    
    # Check if CID exists on IPFS network (if IPFS is available)
    ipfs_verified = False
    ipfs_verification_msg = "IPFS daemon offline - skipped network verification"
    if ipfs_connected and cid and ipfs_http_available:
        ipfs_verified = verify_cid_on_ipfs(cid)
        ipfs_verification_msg = "✓ Verified on IPFS network" if ipfs_verified else "✗ Not found on IPFS network"
    
    # Create chain of custody entry
    verification_status = "Valid CID" if verification_ok else "Invalid CID"
    ipfs_status = "Verified" if ipfs_verified else ("Offline" if not ipfs_connected else "Not Found")
    
    entry = {
        "action": "CID Verified",
        "performedBy": f"{user['name']} ({user['role'].capitalize()})",
        "timestamp": now_iso(),
        "remarks": f"CID: {verification_status}, IPFS: {ipfs_status}",
    }
    
    # Update evidence record
    update_doc = {"$push": {"chainOfCustody": entry}}
    
    if verification_ok:
        update_doc["$set"] = {
            "cid_verified": True,
            "cid_verified_by": f"{user['name']} ({user['role'].capitalize()})",
            "cid_verified_at": now_iso(),
            "ipfs_verified": ipfs_verified,
            "ipfs_verified_at": now_iso() if ipfs_verified else None,
        }
        logger.info(f"✓ CID verified for evidence {evidence_id} by {user['email']}")
    else:
        update_doc["$set"] = {
            "cid_verified": False,
            "ipfs_verified": ipfs_verified,
        }
        logger.warning(f"✗ CID verification failed for evidence {evidence_id}: file_matches={file_matches}, record_matches={record_matches}")
    
    await db.evidence.update_one({"evidence_id": evidence_id}, update_doc)
    await log_activity(
        user["email"],
        "cid_verified" if verification_ok else "cid_verification_failed",
        evidence_id
    )
    
    # Prepare response message
    if verification_ok and ipfs_verified:
        message = f"✅ File Retrieved Successfully - CID verified on IPFS network"
    elif verification_ok:
        message = f"✅ Valid CID - File authentic and matches evidence record (IPFS verification skipped)"
    else:
        reason = []
        if not file_matches:
            reason.append(f"file CID mismatch (generated: {generated_cid[:16]}..., entered: {cid[:16]}...)")
        if not record_matches:
            reason.append(f"record CID mismatch (stored: {ev.get('cid', 'None')[:16]}..., entered: {cid[:16]}...)")
        message = f"❌ Invalid CID - {' and '.join(reason)}"
    
    return {
        "ok": verification_ok,
        "entered_cid": cid,
        "generated_cid": generated_cid,
        "expected_cid": ev.get("cid"),
        "file_matches": file_matches,
        "record_matches": record_matches,
        "ipfs_verified": ipfs_verified,
        "ipfs_status": ipfs_status,
        "message": message,
        "role": user["role"],
        "verified_by": user["name"],
    }



@api.post("/evidence/{evidence_id}/approve")
async def approve_evidence(evidence_id: str, user: dict = Depends(require_roles("investigator", "forensic", "judge", "court_officer"))):
    ev = await db.evidence.find_one({"evidence_id": evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    existing = ev.get("approvals", []) or []
    if any(a.get("user") == user["email"] for a in existing):
        return {"ok": True, "message": "Already approved", "approvals": existing}
    approval = {
        "role": user["role"].capitalize(),
        "user": user["email"],
        "timestamp": now_iso(),
    }
    entry = {
        "action": "Approved",
        "performedBy": f"{user['name']} ({user['role'].capitalize()})",
        "timestamp": now_iso(),
        "remarks": "",
    }
    await db.evidence.update_one(
        {"evidence_id": evidence_id},
        {"$push": {"approvals": approval, "chainOfCustody": entry}},
    )
    await log_activity(user["email"], "evidence_approved", evidence_id)
    return {"ok": True, "message": "Evidence approved", "approvals": existing + [approval]}


@api.get("/evidence/verify/{evidence_id}")
async def verify_evidence_compat(evidence_id: str, user: dict = Depends(get_current_user)):
    return await verify_evidence(evidence_id, user)

@api.get("/evidence/{evidence_id}/download")
async def download_evidence(evidence_id: str, user: dict = Depends(get_current_user)):
    ev = await db.evidence.find_one({"evidence_id": evidence_id})
    if not ev:
        raise HTTPException(404, "Not found")
    path = UPLOAD_DIR / ev["stored_filename"]
    if not path.exists():
        raise HTTPException(404, "File missing")
    return FileResponse(str(path), filename=ev["original_filename"])


@api.get("/evidence/{evidence_id}/ipfs-gateway-url")
async def get_ipfs_gateway_url(evidence_id: str, user: dict = Depends(get_current_user)):
    """
    Get IPFS gateway URL for accessing evidence file from IPFS.
    Returns error if CID not available or IPFS is offline.
    """
    ev = await db.evidence.find_one({"evidence_id": evidence_id})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    
    cid = ev.get("cid")
    if not cid:
        raise HTTPException(
            400,
            detail="Evidence does not have a CID. IPFS may have been offline when file was uploaded."
        )
    
    # Build the IPFS gateway URL
    gateway_url = f"{IPFS_GATEWAY_URL}/{cid}"
    
    return {
        "evidence_id": evidence_id,
        "cid": cid,
        "gateway_url": gateway_url,
        "filename": ev.get("original_filename", "evidence.bin"),
        "ipfs_status": ev.get("ipfs_status", "unknown"),
        "message": "✓ Click link to open from IPFS" if ipfs_connected else "Note: Local IPFS gateway - ensure daemon is running"
    }

# ---------------- Witness Management ----------------
@api.post("/witnesses")
async def create_witness(body: WitnessCreateIn, user: dict = Depends(require_roles("investigator", "police"))):
    """Create a new witness record for a case."""
    witness_id = f"WIT-{str(uuid.uuid4())[:8].upper()}"
    timestamp = now_iso()

    # Verify case exists
    case = await db.cases.find_one({"case_id": body.case_id})
    if not case:
        raise HTTPException(404, "Case not found")

    doc = {
        "witness_id": witness_id,
        "case_id": body.case_id,
        "fir_id": body.fir_id or case.get("fir_id"),
        "name": body.name,
        "contact_info": body.contact_info,
        "statement": body.statement,
        "is_protected": body.is_protected,
        "is_confidential": body.is_confidential,
        "uploaded_by": user["email"],
        "uploaded_at": timestamp,
        "documents": [],
        "verification_status": "pending",
        "verified_by": None,
        "verified_at": None,
        # Normalize naming so UI works consistently.
        "chain_of_custody": [],
    }

    # Seed initial chain entry in the normalized list.
    doc["chain_of_custody"].append({
        "action": "Witness Record Created",
        "performed_by": f"{user['name']} ({user['role'].capitalize()})",
        "timestamp": timestamp,
        "remarks": "Initial witness record"
    })

    
    await db.witnesses.insert_one(doc)
    await append_block(body.case_id, "WITNESS_CREATED", {"witness_id": witness_id}, user["email"])
    await log_activity(user["email"], "witness_create", witness_id)
    
    doc.pop("_id", None)
    return doc


@api.get("/witnesses")
async def list_witnesses(case_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    """List witnesses, optionally filtered by case_id."""
    query = {}
    if case_id:
        query["case_id"] = case_id
    
    # Citizens can only see witnesses for their own cases
    if user["role"] == "citizen":
        user_cases = await db.cases.find({"citizen_id": user["id"]}, {"case_id": 1}).to_list(500)
        user_case_ids = [c["case_id"] for c in user_cases]
        query["case_id"] = {"$in": user_case_ids}
        if case_id and case_id not in user_case_ids:
            raise HTTPException(403, "Access denied")
    
    witnesses = await db.witnesses.find(query, {"_id": 0}).sort("uploaded_at", -1).to_list(500)
    return {"witnesses": witnesses, "count": len(witnesses)}


@api.get("/witnesses/{witness_id}")
async def get_witness(witness_id: str, user: dict = Depends(get_current_user)):
    """Get a specific witness record."""
    witness = await db.witnesses.find_one({"witness_id": witness_id}, {"_id": 0})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    # Citizens can only access witnesses for their own cases
    if user["role"] == "citizen":
        case = await db.cases.find_one({"case_id": witness["case_id"], "citizen_id": user["id"]})
        if not case:
            raise HTTPException(403, "Access denied")
    
    return witness


@api.put("/witnesses/{witness_id}")
async def update_witness(witness_id: str, body: WitnessUpdateIn, user: dict = Depends(require_roles("investigator", "police"))):
    """Update witness information (only before verification)."""
    witness = await db.witnesses.find_one({"witness_id": witness_id})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    if witness.get("verification_status") == "verified":
        raise HTTPException(400, "Cannot modify verified witness record")
    
    update_data = {k: v for k, v in body.model_dump().items() if v is not None}
    if not update_data:
        raise HTTPException(400, "No fields to update")
    
    update_data["updated_by"] = user["email"]
    update_data["updated_at"] = now_iso()
    
    await db.witnesses.update_one(
        {"witness_id": witness_id},
        {"$set": update_data, "$push": {"chain_of_custody": {
            "action": "Witness Updated",
            "performed_by": f"{user['name']} ({user['role'].capitalize()})",
            "timestamp": now_iso(),
            "remarks": f"Updated fields: {', '.join(update_data.keys())}"
        }}}
    )
    
    await append_block(witness["case_id"], "WITNESS_UPDATED", {"witness_id": witness_id}, user["email"])
    await log_activity(user["email"], "witness_update", witness_id)
    
    updated = await db.witnesses.find_one({"witness_id": witness_id}, {"_id": 0})
    return updated


@api.post("/witnesses/{witness_id}/documents")
async def upload_witness_document(
    witness_id: str,
    file: UploadFile = File(...),
    description: str = Form(""),
    user: dict = Depends(require_roles("investigator", "police", "citizen"))
):
    """Upload a document for a witness with SHA-256 hash and IPFS CID."""
    witness = await db.witnesses.find_one({"witness_id": witness_id})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    # Citizens can only upload to their own cases
    if user["role"] == "citizen":
        case = await db.cases.find_one({"case_id": witness["case_id"], "citizen_id": user["id"]})
        if not case:
            raise HTTPException(403, "Access denied")
    
    content = await file.read()
    file_hash = sha256_hex(content)
    
    # Attempt IPFS upload
    cid = None
    ipfs_status = "offline"
    refresh_ipfs_state()
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "witness_doc.bin", content)
        ipfs_status = "online" if cid else "failed"
        logger.info(f"IPFS upload attempt: {ipfs_status} for witness document {file.filename}")
    else:
        logger.warning(f"IPFS daemon offline - storing witness document locally only: {file.filename}")
    
    doc_id = str(uuid.uuid4())
    ext = Path(file.filename or "").suffix
    stored_name = f"witness_{doc_id}{ext}"
    timestamp = now_iso()
    
    # Store file locally
    (UPLOAD_DIR / stored_name).write_bytes(content)
    
    document_record = {
        "document_id": f"WD-{str(uuid.uuid4())[:8].upper()}",
        "witness_id": witness_id,
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "size_bytes": len(content),
        "sha256_hash": file_hash,
        "cid": cid,
        "ipfs_status": ipfs_status,
        "description": description,
        "uploaded_by": user["email"],
        "uploaded_at": timestamp,
        "verification_status": "pending"
    }
    
    await db.witnesses.update_one(
        {"witness_id": witness_id},
        {"$push": {"documents": document_record}, "$set": {"updated_at": timestamp}}
    )
    
    await append_block(witness["case_id"], "WITNESS_DOCUMENT_UPLOADED", {
        "witness_id": witness_id,
        "document_id": document_record["document_id"],
        "hash": file_hash,
        "cid": cid
    }, user["email"])
    
    await log_activity(user["email"], "witness_document_upload", document_record["document_id"])
    
    return document_record


@api.post("/witnesses/{witness_id}/verify")
async def verify_witness(witness_id: str, body: WitnessVerifyIn, user: dict = Depends(require_roles("investigator", "forensic", "court_officer", "judge"))):
    """Verify witness document by SHA-256 hash and/or CID.

    Required consistency with evidence verification:
    - SHA-256 check must be done against stored witness document sha256_hash.
    - CID check must be done against stored witness document cid.
    - Optionally, if IPFS is available, we verify CID existence on the network.

    Returns one of:
    - AUTHENTICATED
    - TAMPERED
    - CID VERIFIED
    - INVALID CID
    """
    witness = await db.witnesses.find_one({"witness_id": witness_id}, {"_id": 0})
    if not witness:
        raise HTTPException(404, "Witness not found")

    if not body.hash and not body.cid:
        raise HTTPException(400, "Provide either hash or cid")

    docs = witness.get("documents", []) or []
    result = {
        "witness_id": witness_id,
        "verified": False,
        "message": "",
        "document_id": None,
        "cid_verified": None,
        "hash_verified": None,
        "ipfs_verified": None,
    }

    # SHA-256 verification
    if body.hash:
        matched = next((d for d in docs if d.get("sha256_hash") == body.hash), None)
        result["hash_verified"] = matched is not None
        if matched:
            result["document_id"] = matched.get("document_id")
            result["message"] = "AUTHENTICATED"
            # Final verified is computed after CID checks if provided.
            result["verified"] = (not body.cid)
        else:
            result["verified"] = False
            result["message"] = "TAMPERED"


    # CID verification
    if body.cid:
        matched = next((d for d in docs if d.get("cid") == body.cid), None)
        result["cid_verified"] = matched is not None

        if matched:
            # Network verification for CID consistency with evidence behavior.
            ipfs_checked = False
            ipfs_ok = False
            if matched.get("cid") and ipfs_connected and ipfs_http_available:
                ipfs_checked = True
                ipfs_ok = verify_cid_on_ipfs(matched.get("cid"))
            result["ipfs_verified"] = ipfs_ok

            cid_record_ok = bool(result.get("cid_verified"))

            # If hash was also provided, require BOTH hash + CID record match.
            if body.hash:
                hash_ok = bool(result.get("hash_verified"))
                result["verified"] = cid_record_ok and hash_ok
                result["message"] = "CID VERIFIED" if result["verified"] else "TAMPERED"
            else:
                # Only CID verification requested: accept when CID matches stored record.
                result["verified"] = cid_record_ok
                if ipfs_checked:
                    result["message"] = "CID VERIFIED" if ipfs_ok else "INVALID CID"
                else:
                    result["message"] = "CID VERIFIED"

            result["document_id"] = result.get("document_id") or matched.get("document_id")

        else:
            result["cid_verified"] = False
            result["verified"] = False
            result["ipfs_verified"] = False
            result["message"] = "INVALID CID"


    # Update witness verification status if hash/CID verified.
    # If both hash+cid were provided, only mark as verified when both are correct.
    if result["verified"]:
        await db.witnesses.update_one(
            {"witness_id": witness_id},
            {
                "$set": {
                    "verification_status": "verified",
                    "verified_by": user["email"],
                    "verified_at": now_iso(),
                }
            },
        )
        await append_block(witness["case_id"], "WITNESS_VERIFIED", {"witness_id": witness_id}, user["email"])
        await log_activity(user["email"], "witness_verify", witness_id)

    return result



@api.post("/witnesses/{witness_id}/forward")
async def forward_witness_to_court(witness_id: str, note: str = Form(""), user: dict = Depends(require_roles("investigator"))):
    """Forward witness record to court."""
    witness = await db.witnesses.find_one({"witness_id": witness_id})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    await db.witnesses.update_one(
        {"witness_id": witness_id},
        {"$set": {"forwarded_to_court": True, "forwarded_at": now_iso(), "forwarded_by": user["email"], "forward_note": note},
         "$push": {"chain_of_custody": {
             "action": "Forwarded to Court",
             "performed_by": f"{user['name']} ({user['role'].capitalize()})",
             "timestamp": now_iso(),
             "remarks": note or "Witness record forwarded to court"
         }}}
    )
    
    await append_block(witness["case_id"], "WITNESS_FORWARDED", {"witness_id": witness_id}, user["email"])
    await log_activity(user["email"], "witness_forward", witness_id)
    
    return {"ok": True, "message": "Witness forwarded to court"}


@api.delete("/witnesses/{witness_id}")
async def delete_witness(witness_id: str, user: dict = Depends(require_roles("investigator", "police"))):
    """Delete a witness record (only if not verified)."""
    witness = await db.witnesses.find_one({"witness_id": witness_id})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    if witness.get("verification_status") == "verified":
        raise HTTPException(400, "Cannot delete verified witness record")
    
    await db.witnesses.delete_one({"witness_id": witness_id})
    await append_block(witness["case_id"], "WITNESS_DELETED", {"witness_id": witness_id}, user["email"])
    await log_activity(user["email"], "witness_delete", witness_id)
    
    return {"ok": True, "message": "Witness deleted"}


@api.get("/witnesses/{witness_id}/download/{document_id}")
async def download_witness_document(witness_id: str, document_id: str, user: dict = Depends(get_current_user)):
    """Download a witness document."""
    witness = await db.witnesses.find_one({"witness_id": witness_id})
    if not witness:
        raise HTTPException(404, "Witness not found")
    
    # Citizens can only download from their own cases
    if user["role"] == "citizen":
        case = await db.cases.find_one({"case_id": witness["case_id"], "citizen_id": user["id"]})
        if not case:
            raise HTTPException(403, "Access denied")
    
    document = None
    for doc in witness.get("documents", []):
        if doc["document_id"] == document_id:
            document = doc
            break
    
    if not document:
        raise HTTPException(404, "Document not found")
    
    path = UPLOAD_DIR / document["stored_filename"]
    if not path.exists():
        raise HTTPException(404, "File not found on server")
    
    return FileResponse(str(path), filename=document["original_filename"])


# ---------------- Forensic ----------------
@api.post("/forensic")
async def add_forensic(body: ForensicCreate, user: dict = Depends(require_roles("forensic"))):
    doc = {
        "id": str(uuid.uuid4()),
        **body.model_dump(),
        "uploaded_by": user["email"],
        "uploaded_at": now_iso(),
    }
    await db.forensic_records.insert_one(doc)
    await append_block(body.case_id, "FORENSIC_ADD", {"case_id": body.case_id}, user["email"])
    await log_activity(user["email"], "forensic_add", body.case_id)
    doc.pop("_id", None)
    return doc

@api.get("/forensic")
async def list_forensic(case_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    if user["role"] == "citizen":
        raise HTTPException(403, "Not allowed")
    query = {"case_id": case_id} if case_id else {}
    return await db.forensic_records.find(query, {"_id": 0}).sort("uploaded_at", -1).to_list(500)

@api.get("/forensic/cases")
async def forensic_cases(user: dict = Depends(require_roles("forensic"))):
    query = {
        "$or": [
            {"status": {"$in": ["Forensic Review", "Sent to Forensic Review"]}},
            {"forwarded_to_forensic": True},
        ]
    }
    cases = await db.cases.find(query, {"_id": 0}).sort("updated_at", -1).to_list(500)
    results = []
    for case_doc in cases:
        evidence = await db.evidence.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("uploaded_at", -1).to_list(300)
        forensic_uploads = await db.forensic_uploads.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("uploaded_at", -1).to_list(300)
        reports = await db.forensic_reports.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("created_at", -1).to_list(100)
        report_files = await db.forensic_report_files.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("uploaded_at", -1).to_list(100)
        results.append({"case": case_doc, "evidence": evidence, "forensic_uploads": forensic_uploads, "reports": reports, "report_files": report_files})
    return results


@api.get("/forensic/report-template")
async def forensic_report_template(user: dict = Depends(require_roles("forensic", "judge", "court_officer"))):
    template_path = ensure_forensic_report_template()
    return FileResponse(str(template_path), filename=template_path.name)


@api.post("/forensic/upload-completed-report")
async def forensic_upload_completed_report(
    case_id: str = Form(...),
    file: UploadFile = File(...),
    description: str = Form(""),
    user: dict = Depends(require_roles("forensic")),
):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty report file")
    report_hash = sha256_hex(content)
    refresh_ipfs_state()
    cid = None
    ipfs_status = "offline"
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "forensic-report.bin", content)
        ipfs_status = "online" if cid else "failed"
    report_file_id = f"RPT-{str(uuid.uuid4())[:6].upper()}"
    ext = Path(file.filename or "").suffix
    stored_name = f"forensic_report_{report_file_id}{ext}"
    ts = now_iso()
    (UPLOAD_DIR / stored_name).write_bytes(content)
    doc = {
        "id": str(uuid.uuid4()),
        "report_file_id": report_file_id,
        "case_id": case_id,
        "fir_id": case_doc.get("fir_id"),
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "size_bytes": len(content),
        "description": (description or "").strip(),
        "sha256_hash": report_hash,
        "cid": cid,
        "ipfs_status": ipfs_status,
        "uploaded_by": user["email"],
        "uploaded_at": ts,
        "approvals": [
            {"role": user["role"], "user": user["email"], "timestamp": ts, "type": "uploaded"}
        ],
        "chainOfCustody": [
            {
                "action": "Forensic Report Uploaded",
                "performedBy": f"{user.get('name', user['email'])} ({user.get('role', 'forensic').capitalize()})",
                "timestamp": ts,
                "remarks": f"IPFS Status: {ipfs_status}",
            }
        ],
    }
    await db.forensic_report_files.insert_one(doc)
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": now_iso()}, "$push": {"forensic_updates": {
            "type": "forensic_completed_report_uploaded",
            "report_file_id": report_file_id,
            "sha256_hash": report_hash,
            "cid": cid,
            "at": ts,
            "by": user["email"],
        }}},
    )
    await append_block(
        case_id,
        "FORENSIC_COMPLETED_REPORT_UPLOADED",
        {"report_file_id": report_file_id, "sha256_hash": report_hash, "cid": cid, "ipfs_status": ipfs_status},
        user["email"],
    )
    await log_activity(user["email"], "forensic_upload_completed_report", f"{case_id}::{report_file_id}")
    doc.pop("_id", None)
    return doc


@api.get("/forensic/reports/{report_file_id}/download")
async def download_forensic_report(report_file_id: str, user: dict = Depends(get_current_user)):
    report = await db.forensic_report_files.find_one({"report_file_id": report_file_id}, {"_id": 0})
    if not report:
        raise HTTPException(404, "Forensic report not found")
    path = UPLOAD_DIR / report["stored_filename"]
    if not path.exists():
        raise HTTPException(404, "Report file missing")
    return FileResponse(str(path), filename=report.get("original_filename") or f"{report_file_id}.bin")


@api.get("/forensic/reports/{report_file_id}/verify-hash")
async def verify_forensic_report_hash(report_file_id: str, user: dict = Depends(require_roles("forensic", "judge", "court_officer"))):
    report = await db.forensic_report_files.find_one({"report_file_id": report_file_id}, {"_id": 0})
    if not report:
        raise HTTPException(404, "Forensic report not found")
    path = UPLOAD_DIR / report["stored_filename"]
    if not path.exists():
        raise HTTPException(404, "Report file missing")
    current_hash = sha256_hex(path.read_bytes())
    ok = current_hash == report.get("sha256_hash")
    await db.forensic_report_files.update_one(
        {"report_file_id": report_file_id},
        {"$push": {
            "approvals": {"role": user["role"], "user": user["email"], "timestamp": now_iso(), "type": "hash_verify", "ok": ok},
            "chainOfCustody": {
                "action": "Forensic Report Hash Verified",
                "performedBy": f"{user.get('name', user['email'])} ({user.get('role', '').capitalize()})",
                "timestamp": now_iso(),
                "remarks": "Hash matched" if ok else "Hash mismatch detected",
            },
        }},
    )
    await append_block(report["case_id"], "FORENSIC_REPORT_HASH_VERIFIED", {"report_file_id": report_file_id, "ok": ok}, user["email"])
    return {"ok": ok, "message": "Hash verified" if ok else "Hash mismatch", "original_hash": report.get("sha256_hash"), "current_hash": current_hash}


@api.get("/forensic/reports/{report_file_id}/verify-cid")
async def verify_forensic_report_cid(report_file_id: str, user: dict = Depends(require_roles("forensic", "judge", "court_officer"))):
    report = await db.forensic_report_files.find_one({"report_file_id": report_file_id}, {"_id": 0})
    if not report:
        raise HTTPException(404, "Forensic report not found")
    cid = report.get("cid")
    if not cid:
        return {"ok": False, "message": "CID not available for this report", "cid": None}
    refresh_ipfs_state()
    ipfs_ok = verify_cid_on_ipfs(cid) if ipfs_connected and ipfs_http_available else False
    await db.forensic_report_files.update_one(
        {"report_file_id": report_file_id},
        {"$set": {"cid_verified": ipfs_ok, "cid_verified_by": user["email"], "cid_verified_at": now_iso()}, "$push": {
            "approvals": {"role": user["role"], "user": user["email"], "timestamp": now_iso(), "type": "cid_verify", "ok": ipfs_ok},
            "chainOfCustody": {
                "action": "Forensic Report CID Verified",
                "performedBy": f"{user.get('name', user['email'])} ({user.get('role', '').capitalize()})",
                "timestamp": now_iso(),
                "remarks": "CID verified on IPFS" if ipfs_ok else "CID verification failed/offline",
            },
        }},
    )
    await append_block(report["case_id"], "FORENSIC_REPORT_CID_VERIFIED", {"report_file_id": report_file_id, "ok": ipfs_ok, "cid": cid}, user["email"])
    return {"ok": ipfs_ok, "message": "CID verified on IPFS" if ipfs_ok else "CID verification failed or IPFS offline", "cid": cid}


@api.post("/forensic/upload-file")
async def forensic_upload_file(
    case_id: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    user: dict = Depends(require_roles("forensic")),
):
    """
    Upload a forensic artifact (file) for a case.
    Always stores locally; attempts IPFS CID generation if daemon is available.
    """
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")

    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file")

    file_hash = sha256_hex(content)
    refresh_ipfs_state()
    cid = None
    ipfs_status = "offline"
    if ipfs_connected and ipfs_http_available:
        cid = upload_to_ipfs(file.filename or "forensic.bin", content)
        ipfs_status = "online" if cid else "failed"

    upload_id = str(uuid.uuid4())
    ext = Path(file.filename or "").suffix
    stored_name = f"forensic_{upload_id}{ext}"
    timestamp = now_iso()
    (UPLOAD_DIR / stored_name).write_bytes(content)

    doc = {
        "id": upload_id,
        "upload_id": f"FRN-{str(uuid.uuid4())[:6].upper()}",
        "case_id": case_id,
        "description": (description or "").strip(),
        "original_filename": file.filename,
        "stored_filename": stored_name,
        "size_bytes": len(content),
        "sha256_hash": file_hash,
        "cid": cid,
        "ipfs_status": ipfs_status,
        "uploaded_by": user["email"],
        "uploaded_at": timestamp,
        "chainOfCustody": [
            {
                "action": "Forensic Artifact Uploaded",
                "performedBy": f"{user.get('name', user['email'])} ({user.get('role', 'forensic').capitalize()})",
                "timestamp": timestamp,
                "remarks": f"IPFS Status: {ipfs_status}",
            }
        ],
    }

    await db.forensic_uploads.insert_one(doc)
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": now_iso()}, "$push": {"forensic_updates": {
            "type": "forensic_file_uploaded",
            "upload_id": doc["upload_id"],
            "filename": doc["original_filename"],
            "sha256_hash": file_hash,
            "cid": cid,
            "ipfs_status": ipfs_status,
            "at": timestamp,
            "by": user["email"],
        }}},
    )
    await append_block(
        case_id,
        "FORENSIC_FILE_UPLOADED",
        {"upload_id": doc["upload_id"], "filename": doc["original_filename"], "sha256_hash": file_hash, "cid": cid, "ipfs_status": ipfs_status},
        user["email"],
    )
    await log_activity(user["email"], "forensic_upload_file", f"{case_id}::{doc['upload_id']}")
    doc.pop("_id", None)
    return doc

@api.post("/forensic/upload-report")
async def forensic_upload_report(body: ForensicReportUploadIn, user: dict = Depends(require_roles("forensic"))):
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Forensic Review Completed")
    report_payload = body.model_dump()
    report_hash = sha256_hex(json.dumps(report_payload, sort_keys=True).encode("utf-8"))
    refresh_ipfs_state()
    report_cid = None
    report_ipfs_status = "offline"
    if ipfs_connected and ipfs_http_available:
        report_cid = upload_to_ipfs(f"forensic-report-{body.case_id}.json", json.dumps(report_payload, sort_keys=True).encode("utf-8"))
        report_ipfs_status = "online" if report_cid else "failed"
    report_doc = {
        "id": str(uuid.uuid4()),
        "case_id": body.case_id,
        "report_title": body.report_title,
        "summary": body.summary,
        "examiner_name": body.examiner_name,
        "lab_name": body.lab_name,
        "verified_evidence_ids": body.verified_evidence_ids or [],
        "matching_hashes": body.matching_hashes or [],
        "ai_tampering_analysis": body.ai_tampering_analysis,
        "final_conclusion": body.final_conclusion,
        "digital_signature": body.digital_signature,
        "image_analysis": body.image_analysis,
        "log_analysis": body.log_analysis,
        "ip_analysis": body.ip_analysis,
        "call_data_analysis": body.call_data_analysis,
        "result": body.result,
        "sha256_hash": report_hash,
        "cid": report_cid,
        "ipfs_status": report_ipfs_status,
        "created_by": user["email"],
        "created_at": now_iso(),
    }
    await db.forensic_reports.insert_one(report_doc)
    status_note = f"Forensic report uploaded: {body.report_title}"
    case_updates = {
        "$set": {"status": "Forensic Review Completed", "updated_at": now_iso()},
        "$push": {
            "status_history": {
                "status": "Forensic Review Completed",
                "at": now_iso(),
                "by": user["email"],
                "note": status_note,
            },
            "forensic_updates": {
                "type": "report_uploaded",
                "report_id": report_doc["id"],
                "result": body.result,
                "at": now_iso(),
                "by": user["email"],
            },
        },
    }
    if body.send_to_court:
        case_updates["$set"]["forensic_result_sent_to_court"] = True
        case_updates["$set"]["forensic_result_sent_to_court_at"] = now_iso()
        case_updates["$set"]["forensic_result_sent_to_court_by"] = user["email"]
        case_updates["$push"]["forensic_updates"]["court_dispatch"] = True
    await db.cases.update_one({"case_id": body.case_id}, case_updates)
    await append_block(
        body.case_id,
        "FORENSIC_REPORT_UPLOADED",
        {"report_title": body.report_title, "result": body.result, "sent_to_court": body.send_to_court, "sha256_hash": report_hash, "cid": report_cid},
        user["email"],
    )
    await log_activity(user["email"], "forensic_upload_report", body.case_id)
    report_doc.pop("_id", None)
    return report_doc

@api.post("/forensic/verify-evidence")
async def forensic_verify_evidence(body: ForensicVerifyEvidenceIn, user: dict = Depends(require_roles("forensic"))):
    if body.verdict not in ("Verified", "Suspicious"):
        raise HTTPException(400, "Invalid verdict")
    ev = await db.evidence.find_one({"evidence_id": body.evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    path = UPLOAD_DIR / ev["stored_filename"]
    if not path.exists():
        raise HTTPException(404, "Evidence file missing")
    recalculated_hash = sha256_hex(path.read_bytes())
    hash_ok = recalculated_hash == ev["sha256_hash"]
    tampered = (not hash_ok) or body.verdict == "Suspicious"
    verification_doc = {
        "id": str(uuid.uuid4()),
        "evidence_id": body.evidence_id,
        "case_id": ev["case_id"],
        "verdict": body.verdict,
        "note": (body.note or "").strip(),
        "hash_ok": hash_ok,
        "stored_hash": ev["sha256_hash"],
        "recalculated_hash": recalculated_hash,
        "tampered": tampered,
        "verified_by": user["email"],
        "verified_at": now_iso(),
    }
    await db.evidence_verifications.insert_one(verification_doc)
    await db.evidence.update_one(
        {"evidence_id": body.evidence_id},
        {"$set": {
            "tampered": tampered,
            "forensic_verdict": body.verdict,
            "forensic_note": (body.note or "").strip(),
            "verified_at": now_iso(),
            "verified_by": user["email"],
        }},
    )
    await append_block(
        ev["case_id"],
        "FORENSIC_EVIDENCE_VERIFIED",
        {"evidence_id": body.evidence_id, "verdict": body.verdict, "tampered": tampered},
        user["email"],
    )
    await log_activity(user["email"], "forensic_verify_evidence", body.evidence_id)
    verification_doc.pop("_id", None)
    return verification_doc

# ---------------- Court Module ----------------
@api.get("/court/cases")
async def court_cases(user: dict = Depends(require_roles("court_officer", "judge"))):
    query = {
        "$or": [
            {"status": "Sent to Court"},
            {"status": "Hearing Scheduled"},
            {"status": "REOPENED"},
            {"hearing_status": {"$in": ["Scheduled", "In Progress", "Completed"]}},
            {"forensic_result_sent_to_court": True},
            {"appeal_status": "Pending"},
        ]
    }
    return await db.cases.find(query, {"_id": 0}).sort("updated_at", -1).to_list(500)

@api.post("/court/schedule-hearing")
async def court_schedule_hearing(body: CourtScheduleHearingIn, user: dict = Depends(require_roles("court_officer", "judge"))):
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Sent to Court")
    ensure_valid_case_transition("Sent to Court", "Hearing Scheduled")
    note = (body.note or "").strip()
    await db.cases.update_one(
        {"case_id": body.case_id},
        {
            "$set": {
                "status": "Hearing Scheduled",
                "hearing_date": body.hearing_date,
                "hearing_status": "Scheduled",
                "updated_at": now_iso(),
            },
            "$push": {
                "status_history": {
                    "status": "Sent to Court",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": "Moved to court hearing queue",
                }
            },
        },
    )
    await db.cases.update_one(
        {"case_id": body.case_id},
        {
            "$push": {
                "status_history": {
                    "status": "Hearing Scheduled",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": note or f"Hearing date set: {body.hearing_date}",
                },
                "court_updates": {
                    "type": "hearing_scheduled",
                    "hearing_date": body.hearing_date,
                    "hearing_status": "Scheduled",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": note,
                },
            }
        },
    )
    await append_block(body.case_id, "COURT_HEARING_SCHEDULED", {"hearing_date": body.hearing_date}, user["email"])
    await log_activity(user["email"], "court_schedule_hearing", body.case_id)
    await sync_fir_status_with_case(body.case_id)
    updated_case = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    schedule_case_notifications(notify_hearing_scheduled(db, updated_case, body.hearing_date, note))
    return updated_case

@api.post("/court/update-status")
async def court_update_status(body: CourtStatusUpdateIn, user: dict = Depends(require_roles("court_officer", "judge"))):
    if body.hearing_status not in ("Scheduled", "In Progress", "Completed"):
        raise HTTPException(400, "Invalid hearing status")
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Hearing Scheduled")
    await db.cases.update_one(
        {"case_id": body.case_id},
        {
            "$set": {
                "hearing_status": body.hearing_status,
                "status": "Hearing Scheduled",
                "updated_at": now_iso(),
            },
            "$push": {
                "court_updates": {
                    "type": "hearing_status_update",
                    "hearing_status": body.hearing_status,
                    "at": now_iso(),
                    "by": user["email"],
                    "note": (body.note or "").strip(),
                }
            },
        },
    )
    await append_block(body.case_id, "COURT_STATUS_UPDATED", {"hearing_status": body.hearing_status}, user["email"])
    await log_activity(user["email"], "court_update_status", f"{body.case_id} → {body.hearing_status}")
    await sync_fir_status_with_case(body.case_id)
    return await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})

# ---------------- Judge Module ----------------
@api.get("/judge/cases")
async def judge_cases(user: dict = Depends(require_roles("judge"))):
    query = {
        "$or": [
            {"status": {"$in": ["Hearing Scheduled", "Judgment Issued", "REOPENED"]}},
            {"hearing_status": "Completed"},
            {"appeal_status": "Pending"},
        ]
    }
    cases = await db.cases.find(query, {"_id": 0}).sort("updated_at", -1).to_list(500)
    results = []
    for case_doc in cases:
        fir = await db.firs.find_one({"fir_id": case_doc.get("fir_id")}, {"_id": 0})
        investigation_notes = await db.investigation_notes.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("created_at", -1).to_list(200)
        forensic_reports = await db.forensic_reports.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("created_at", -1).to_list(100)
        forensic_uploads = await db.forensic_uploads.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("uploaded_at", -1).to_list(100)
        forensic_report_files = await db.forensic_report_files.find(
            {"case_id": case_doc["case_id"]},
            {"_id": 0},
        ).sort("uploaded_at", -1).to_list(100)
        evidence = await db.evidence.find({"case_id": case_doc["case_id"]}, {"_id": 0}).sort("uploaded_at", -1).to_list(500)
        suspects = await db.suspects.find(
            {"associated_cases": {"$in": [case_doc["case_id"]]}},
            {"_id": 0},
        ).to_list(200)
        judgments = await db.judgments.find({"case_id": case_doc["case_id"]}, {"_id": 0}).sort("issued_at", -1).to_list(20)
        results.append({
            "case": case_doc,
            "fir": fir,
            "investigation_notes": investigation_notes,
            "forensic_reports": forensic_reports,
            "forensic_uploads": forensic_uploads,
            "forensic_report_files": forensic_report_files,
            "evidence": evidence,
            "suspects": suspects,
            "judgments": judgments,
        })
    return results

@api.post("/judge/evidence-action")
async def judge_evidence_action(body: JudgeEvidenceActionIn, user: dict = Depends(require_roles("judge"))):
    ev = await db.evidence.find_one({"evidence_id": body.evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")
    action = (body.action or "").strip().lower()
    note = (body.note or "").strip()
    if action == "verify_hash":
        return await verify_evidence(body.evidence_id, user)
    if action == "verify_cid":
        if not ev.get("cid"):
            raise HTTPException(400, "CID not available for this evidence")
        await db.evidence.update_one(
            {"evidence_id": body.evidence_id},
            {"$set": {"cid_verified": True, "cid_verified_by": user["email"], "cid_verified_at": now_iso()}, "$push": {"chainOfCustody": {
                "action": "CID Verified",
                "performedBy": f"{user.get('name', user['email'])} (Judge)",
                "timestamp": now_iso(),
                "remarks": note or "CID verified by judge",
            }}},
        )
        await append_block(ev["case_id"], "JUDGE_EVIDENCE_CID_VERIFIED", {"evidence_id": body.evidence_id}, user["email"])
        await log_activity(user["email"], "judge_evidence_cid_verify", body.evidence_id)
        return {"ok": True, "message": "CID verified"}
    if action == "approve":
        return await approve_evidence(body.evidence_id, user)
    if action == "reject":
        await db.evidence.update_one(
            {"evidence_id": body.evidence_id},
            {"$set": {"judge_rejected": True, "judge_rejected_note": note, "judge_rejected_at": now_iso(), "judge_rejected_by": user["email"]}, "$push": {"chainOfCustody": {
                "action": "Rejected",
                "performedBy": f"{user.get('name', user['email'])} (Judge)",
                "timestamp": now_iso(),
                "remarks": note or "Rejected by judge",
            }}},
        )
        await append_block(ev["case_id"], "JUDGE_EVIDENCE_REJECTED", {"evidence_id": body.evidence_id, "note": note}, user["email"])
        await log_activity(user["email"], "judge_evidence_reject", body.evidence_id)
        return {"ok": True, "message": "Evidence rejected and logged"}
    raise HTTPException(400, "Invalid action")

@api.post("/judge/return-case")
async def judge_return_case(body: JudgeCaseReturnIn, user: dict = Depends(require_roles("judge"))):
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    target = (body.target or "").strip().lower()
    if target not in ("investigator", "forensic"):
        raise HTTPException(400, "Invalid return target")
    new_status = "Under Investigation" if target == "investigator" else "Forensic Review"
    await db.cases.update_one(
        {"case_id": body.case_id},
        {"$set": {"status": new_status, "updated_at": now_iso()}, "$push": {"status_history": {
            "status": new_status,
            "at": now_iso(),
            "by": user["email"],
            "note": f"Returned to {target}: {(body.note or '').strip()}",
        }}},
    )
    await append_block(body.case_id, "JUDGE_CASE_RETURNED", {"target": target, "note": (body.note or "").strip()}, user["email"])
    await log_activity(user["email"], "judge_return_case", f"{body.case_id}::{target}")
    await sync_fir_status_with_case(body.case_id)
    return {"ok": True, "case_id": body.case_id, "target": target, "status": new_status}

@api.post("/judge/submit-verdict")
async def judge_submit_verdict(
    case_id: str = Form(...),
    verdict: str = Form(...),  # Guilty, Not Guilty, Further Investigation
    decision_note: str = Form(""),
    accused_suspect_id: str = Form(""),
    sentence: str = Form(""),
    sentence_duration_years: int = Form(0),
    sentence_duration_months: int = Form(0),
    fine_amount: float = Form(0.0),
    parole_eligible: bool = Form(False),
    imprisonment_start_date: str = Form(""),
    ppc_sections: str = Form(""),
    hearing_notes: str = Form(""),
    judge_remarks: str = Form(""),
    final_order_file: Optional[UploadFile] = File(None),
    user: dict = Depends(require_roles("judge")),
):
    if verdict not in ("Guilty", "Not Guilty", "Further Investigation"):
        raise HTTPException(400, "Invalid verdict")
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Judgment Issued")
    if case_doc.get("hearing_status") != "Completed":
        raise HTTPException(400, "Invalid case state transition")
    order_meta = None
    if final_order_file is not None:
        content = await final_order_file.read()
        oid = str(uuid.uuid4())
        ext = Path(final_order_file.filename or "").suffix
        stored_name = f"order-{oid}{ext}"
        (UPLOAD_DIR / stored_name).write_bytes(content)
        order_meta = {
            "id": oid,
            "original_filename": final_order_file.filename,
            "stored_filename": stored_name,
            "size_bytes": len(content),
            "sha256_hash": sha256_hex(content),
        }
        await db.final_orders.insert_one({
            "id": oid,
            "case_id": case_id,
            "uploaded_by": user["email"],
            "uploaded_at": now_iso(),
            **order_meta,
        })
    judgment_doc = {
        "id": str(uuid.uuid4()),
        "case_id": case_id,
        "accused_suspect_id": (accused_suspect_id or "").strip() or None,
        "verdict": verdict,
        "decision_note": (decision_note or "").strip(),
        "sentence": (sentence or "").strip(),
        "sentence_duration_years": sentence_duration_years,
        "sentence_duration_months": sentence_duration_months,
        "fine_amount": fine_amount,
        "parole_eligible": parole_eligible,
        "imprisonment_start_date": (imprisonment_start_date or "").strip() or None,
        "expected_release_date": None,
        "ppc_sections": [s.strip() for s in (ppc_sections or "").split(",") if s.strip()],
        "hearing_notes": (hearing_notes or "").strip(),
        "judge_remarks": (judge_remarks or "").strip(),
        "order_file": order_meta,
        "issued_by": user["email"],
        "issued_at": now_iso(),
    }

    # Calculate expected release date if guilty and imprisonment details provided
    if verdict == "Guilty" and (sentence_duration_years > 0 or sentence_duration_months > 0) and imprisonment_start_date:
        try:
            start_date = datetime.fromisoformat(imprisonment_start_date.replace('Z', '+00:00'))
            total_months = sentence_duration_years * 12 + sentence_duration_months
            release_date = start_date + timedelta(days=total_months * 30)  # Approximate months to days
            judgment_doc["expected_release_date"] = release_date.isoformat()
        except Exception:
            pass  # Invalid date format, leave as None
    await db.judgments.insert_one(judgment_doc)
    await db.cases.update_one(
        {"case_id": case_id},
        {
            "$set": {
                "status": "Judgment Issued",
                "judgment_verdict": verdict,
                "judgment_note": (decision_note or "").strip(),
                "judgment_accused_suspect_id": (accused_suspect_id or "").strip() or None,
                "judgment_sentence": (sentence or "").strip(),
                "judgment_sentence_duration_years": sentence_duration_years,
                "judgment_sentence_duration_months": sentence_duration_months,
                "judgment_fine_amount": fine_amount,
                "judgment_parole_eligible": parole_eligible,
                "judgment_imprisonment_start_date": (imprisonment_start_date or "").strip() or None,
                "judgment_expected_release_date": judgment_doc["expected_release_date"],
                "judgment_ppc_sections": [s.strip() for s in (ppc_sections or "").split(",") if s.strip()],
                "judgment_hearing_notes": (hearing_notes or "").strip(),
                "judgment_remarks": (judge_remarks or "").strip(),
                "updated_at": now_iso(),
            },
            "$push": {
                "status_history": {
                    "status": "Judgment Issued",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": f"Verdict: {verdict}",
                },
                "judge_updates": {
                    "type": "verdict_submitted",
                    "verdict": verdict,
                    "at": now_iso(),
                    "by": user["email"],
                    "note": (decision_note or "").strip(),
                },
            },
        },
    )

    # Update suspect sentence status if guilty verdict with accused
    if verdict == "Guilty" and accused_suspect_id:
        sentence_status = "Active Sentence" if (sentence_duration_years > 0 or sentence_duration_months > 0) else "Completed Sentence"
        suspect_update = {
            "sentence_status": sentence_status,
            "sentence_duration_years": sentence_duration_years,
            "sentence_duration_months": sentence_duration_months,
            "fine_amount": fine_amount,
            "parole_eligible": parole_eligible,
            "imprisonment_start_date": (imprisonment_start_date or "").strip() or None,
            "expected_release_date": judgment_doc["expected_release_date"],
            "case_id": case_id,
            "verdict_issued_at": now_iso(),
            "updated_at": now_iso(),
        }
        await db.suspects.update_one(
            {"suspect_id": accused_suspect_id},
            {"$set": suspect_update}
        )

    await append_block(case_id, "JUDGMENT_ISSUED", {"verdict": verdict}, user["email"])
    await log_activity(user["email"], "judge_submit_verdict", f"{case_id} → {verdict}")
    await sync_fir_status_with_case(case_id)
    judgment_doc.pop("_id", None)
    updated_case = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    schedule_case_notifications(notify_verdict_issued(db, updated_case, judgment_doc))
    return judgment_doc

@api.post("/judge/close-case")
async def judge_close_case(body: JudgeCloseCaseIn, user: dict = Depends(require_roles("judge"))):
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    ensure_valid_case_transition(case_doc.get("status"), "Closed")
    await db.cases.update_one(
        {"case_id": body.case_id},
        {
            "$set": {"status": "Closed", "closed_at": now_iso(), "closed_by": user["email"], "updated_at": now_iso()},
            "$push": {
                "status_history": {
                    "status": "Closed",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": (body.note or "").strip(),
                }
            },
        },
    )
    await append_block(body.case_id, "CASE_CLOSED", {"note": (body.note or "").strip()}, user["email"])
    await log_activity(user["email"], "judge_close_case", body.case_id)
    await sync_fir_status_with_case(body.case_id)
    updated_case = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    schedule_case_notifications(notify_case_closed(db, updated_case, (body.note or "").strip()))
    return updated_case


@api.post("/cases/{case_id}/appeal")
async def request_appeal(
    case_id: str,
    reason: str = Form(...),
    lawyer_notes: str = Form(""),
    files: Optional[List[UploadFile]] = File(None),
    user: dict = Depends(require_roles("citizen")),
):
    reason = (reason or "").strip()
    if not reason:
        raise HTTPException(400, "Appeal reason is required")
    lawyer_notes = (lawyer_notes or "").strip()
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    if case_doc.get("citizen_id") != user["id"]:
        raise HTTPException(403, "Not allowed")
    if case_doc.get("status") != "Closed":
        raise HTTPException(400, "Higher court appeal may only be filed after the case is closed")
    existing = await db.appeals.find_one({"case_id": case_id, "status": "Pending"})
    if existing:
        raise HTTPException(400, "An appeal is already pending for this case")
    file_list = list(files or [])
    if len(file_list) > 15:
        raise HTTPException(400, "Too many supporting files (max 15)")
    documents = []
    refresh_ipfs_state()
    for file in file_list:
        if not file or not file.filename:
            continue
        content = await file.read()
        if len(content) > 20 * 1024 * 1024:
            raise HTTPException(400, f"File too large: {file.filename}")
        sha256_hash = hashlib.sha256(content).hexdigest()
        cid = None
        if ipfs_connected and ipfs_http_available:
            try:
                cid = upload_to_ipfs(file.filename or "appeal-doc.bin", content)
            except Exception as e:
                logger.warning(f"IPFS upload failed for appeal doc {file.filename}: {e}")
        stored_name = f"APL-{case_id.replace('/', '-')}-{uuid.uuid4().hex[:10]}_{file.filename}"
        (UPLOAD_DIR / stored_name).write_bytes(content)
        documents.append({
            "filename": file.filename,
            "file_type": file.content_type or "application/octet-stream",
            "sha256_hash": sha256_hash,
            "cid": cid,
            "stored_filename": stored_name,
            "uploaded_at": now_iso(),
            "uploaded_by": user["email"],
        })
    appeal_doc = {
        "id": str(uuid.uuid4()),
        "appeal_id": f"APL-{str(uuid.uuid4())[:6].upper()}",
        "case_id": case_id,
        "fir_id": case_doc.get("fir_id"),
        "requested_by": user["email"],
        "requested_at": now_iso(),
        "reason": reason,
        "lawyer_notes": lawyer_notes,
        "appeal_type": "higher_court",
        "documents": documents,
        "status": "Pending",
        "decision_note": None,
        "decided_by": None,
        "decided_at": None,
    }
    await db.appeals.insert_one(appeal_doc)
    await db.cases.update_one(
        {"case_id": case_id},
        {
            "$set": {"appeal_status": "Pending", "updated_at": now_iso()},
            "$push": {
                "status_history": {
                    "status": "Appeal Pending",
                    "at": now_iso(),
                    "by": user["email"],
                    "note": f"Higher court appeal: {reason[:400]}",
                },
                "appeal_timeline": {
                    "at": now_iso(),
                    "type": "higher_court_filed",
                    "appeal_id": appeal_doc["appeal_id"],
                    "by": user["email"],
                    "summary": reason[:500],
                    "documents_count": len(documents),
                },
            },
        },
    )
    block_data = {
        "appeal_id": appeal_doc["appeal_id"],
        "appeal_type": "higher_court",
        "reason_preview": reason[:240],
        "lawyer_notes_preview": lawyer_notes[:120],
        "document_count": len(documents),
        "document_hashes": [d["sha256_hash"] for d in documents][:30],
    }
    await append_block(case_id, "HIGHER_COURT_APPEAL_FILED", block_data, user["email"])
    await log_activity(user["email"], "appeal_requested", f"{case_id}::{appeal_doc['appeal_id']}")
    appeal_doc.pop("_id", None)
    return appeal_doc


@api.get("/appeals")
async def list_appeals(case_id: Optional[str] = None, user: dict = Depends(get_current_user)):
    query = {}
    if case_id:
        query["case_id"] = case_id
    if user["role"] == "citizen":
        my_firs = await db.firs.find({"citizen_id": user["id"]}, {"fir_id": 1, "_id": 0}).to_list(500)
        my_ids = [f["fir_id"] for f in my_firs]
        query["fir_id"] = {"$in": my_ids}
    elif user["role"] in ("judge", "court_officer"):
        query = query
    elif user["role"] == "admin":
        query = query
    else:
        raise HTTPException(403, "Not allowed")
    appeals = await db.appeals.find(query, {"_id": 0}).sort("requested_at", -1).to_list(500)
    return appeals


@api.post("/cases/{case_id}/appeal/decision")
async def decide_appeal(case_id: str, body: AppealDecisionIn, user: dict = Depends(require_roles("judge", "court_officer"))):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    appeal_doc = await db.appeals.find_one({"case_id": case_id, "status": "Pending"}, {"_id": 0})
    if not appeal_doc:
        raise HTTPException(404, "Pending appeal not found")
    decision = (body.decision or "").strip().lower()
    if decision not in ("accept", "reject"):
        raise HTTPException(400, "Invalid decision")
    update_payload = {
        "$set": {
            "status": "Accepted" if decision == "accept" else "Rejected",
            "decision_note": (body.note or "").strip(),
            "decided_by": user["email"],
            "decided_at": now_iso(),
        }
    }
    await db.appeals.update_one({"case_id": case_id, "status": "Pending"}, update_payload)
    note_dec = (body.note or "").strip()
    if decision == "accept":
        await db.cases.update_one(
            {"case_id": case_id},
            {
                "$set": {"status": "REOPENED", "appeal_status": "Accepted", "updated_at": now_iso()},
                "$push": {
                    "status_history": {
                        "status": "REOPENED",
                        "at": now_iso(),
                        "by": user["email"],
                        "note": note_dec or "Higher court appeal accepted — case reopened",
                    },
                    "appeal_timeline": {
                        "at": now_iso(),
                        "type": "higher_court_accepted",
                        "appeal_id": appeal_doc.get("appeal_id"),
                        "by": user["email"],
                        "summary": note_dec[:500],
                    },
                },
            },
        )
        await append_block(
            case_id,
            "APPEAL_ACCEPTED",
            {
                "note": note_dec,
                "appeal_id": appeal_doc.get("appeal_id"),
                "new_status": "REOPENED",
                "appeal_type": appeal_doc.get("appeal_type", "higher_court"),
            },
            user["email"],
        )
        await log_activity(user["email"], "appeal_accepted", case_id)
        await sync_fir_status_with_case(case_id)
    else:
        await db.cases.update_one(
            {"case_id": case_id},
            {
                "$set": {"appeal_status": "Rejected", "updated_at": now_iso()},
                "$push": {
                    "appeal_timeline": {
                        "at": now_iso(),
                        "type": "higher_court_rejected",
                        "appeal_id": appeal_doc.get("appeal_id"),
                        "by": user["email"],
                        "summary": note_dec[:500],
                    }
                },
            },
        )
        await append_block(
            case_id,
            "APPEAL_REJECTED",
            {
                "note": note_dec,
                "appeal_id": appeal_doc.get("appeal_id"),
                "case_status_preserved": "Closed",
            },
            user["email"],
        )
        await log_activity(user["email"], "appeal_rejected", case_id)
    updated_case = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if decision == "accept":
        schedule_case_notifications(
            notify_higher_court_reopen_broadcast(db, updated_case, appeal_doc.get("appeal_id"))
        )
    schedule_case_notifications(
        notify_appeal_decision(db, updated_case, appeal_doc, decision, note_dec)
    )
    return {"ok": True, "case_id": case_id, "decision": decision}

# ---------------- Blockchain ----------------
@api.get("/blockchain")
async def list_blockchain(user: dict = Depends(get_current_user)):
    blocks = await db.blockchain.find({}, {"_id": 0}).sort("index", 1).to_list(1000)
    # verify integrity while preserving legacy chain compatibility
    integrity_ok = True
    legacy_blocks = 0
    valid_blocks = 0
    broken_blocks = 0
    broken_at = None

    for i, b in enumerate(blocks):
        previous_hash = _get_previous_hash_from_block(b)
        current_hash = _get_current_hash_from_block(b)
        action = b.get("action") or b.get("action_type") or ""
        case_id = b.get("case_id") or b.get("caseId") or ""
        timestamp = b.get("timestamp") or ""
        data = _normalize_block_data(b.get("data") or {})

        recomputed = _generate_block_hash(case_id, action, data, timestamp, previous_hash)
        hash_ok = recomputed == current_hash
        legacy = False

        # Legacy compatibility for older blocks hashed with older concatenation formats.
        if not hash_ok:
            legacy_candidates = []
            legacy_candidates.append(
                sha256_hex(
                    f"{b.get('index')}|{timestamp}|{case_id}|{action}|{str(b.get('data'))}|{previous_hash}|{b.get('user_email')}|{b.get('wallet_id')}".encode()
                )
            )
            legacy_candidates.append(
                sha256_hex(
                    f"{b.get('index')}|{timestamp}|{case_id}|{action}|{str(b.get('data'))}|{previous_hash}|{b.get('user_email')}".encode()
                )
            )
            if current_hash in legacy_candidates:
                legacy_blocks += 1
                legacy = True
                hash_ok = True
            else:
                hash_ok = False

        link_ok = True
        if i > 0 and _get_current_hash_from_block(blocks[i - 1]) != previous_hash:
            link_ok = False

        valid = hash_ok and link_ok
        if valid:
            valid_blocks += 1
        else:
            broken_blocks += 1
            integrity_ok = False
            if broken_at is None:
                broken_at = b.get("index")

        # Ensure clients always receive standardized blockchain fields.
        b["previous_hash"] = previous_hash
        b["previousHash"] = previous_hash
        b["current_hash"] = current_hash
        b["hash"] = current_hash
        b["data"] = data
        b["hash_ok"] = hash_ok
        b["link_ok"] = link_ok
        b["valid"] = valid
        b["legacy"] = legacy

    summary = {
        "integrity_ok": integrity_ok,
        "integrity_status": "legacy_data" if integrity_ok and legacy_blocks > 0 else ("valid" if integrity_ok else "broken"),
        "legacy_blocks": legacy_blocks,
        "valid_blocks": valid_blocks,
        "broken_blocks": broken_blocks,
        "broken_at": broken_at,
        "count": len(blocks),
    }

    if user["role"] == "citizen":
        summary["blocks"] = []
        summary["public_verification"] = {
            "message": "Read-only blockchain audit access provided for citizen accounts.",
            "details": "Contact a court officer or investigator for a full audit report.",
        }
        return summary

    summary["blocks"] = blocks
    return summary

# ---------------- Activity Logs ----------------
@api.get("/activity-logs")
async def list_activity(user: dict = Depends(get_current_user)):
    if user["role"] not in ("admin", "police", "judge"):
        raise HTTPException(403, "Not allowed")
    return await db.activity_logs.find({}, {"_id": 0}).sort("timestamp", -1).limit(500).to_list(500)

# ---------------- Analytics ----------------
@api.post("/ai/predict-risk")
async def ai_predict_risk(body: dict, user: dict = Depends(get_current_user)):
    city = (body.get("city") or "").strip()
    area = (body.get("area") or "").strip()
    if not city or not area:
        raise HTTPException(400, "City and area are required")
    return _predict_crime_risk(city, area)


@api.get("/ai/high-risk-areas")
async def ai_high_risk_areas(city: Optional[str] = None, user: dict = Depends(get_current_user)):
    return {"areas": _get_high_risk_areas(city or "")}


class AIAskIn(BaseModel):
    question: str
    city: Optional[str] = ""
    area: Optional[str] = ""


class HashVerificationIn(BaseModel):
    hash: str

class CIDVerificationIn(BaseModel):
    cid: str

class SmartContractValidationIn(BaseModel):
    evidence_id: str = Field(min_length=1)
    approval_count: int = Field(ge=0)
    hash_value: str = Field(min_length=1)

@api.post("/verify/hash")
async def verify_hash(body: HashVerificationIn, user: dict = Depends(require_roles("investigator", "forensic", "court_officer", "judge", "admin"))):
    # Search for hash in FIR documents and evidence
    fir_doc = await db.firs.find_one({"documents.sha256_hash": body.hash}, {"_id": 0})
    evidence_doc = await db.evidence.find_one({"sha256_hash": body.hash}, {"_id": 0})

    if fir_doc:
        doc = next((d for d in fir_doc.get("documents", []) if d.get("sha256_hash") == body.hash), None)
        if doc:
            return {
                "verified": True,
                "type": "fir_document",
                "fir_id": fir_doc["fir_id"],
                "filename": doc["filename"],
                "uploaded_at": doc["uploaded_at"],
                "uploaded_by": doc["uploaded_by"],
                "cid": doc.get("cid"),
                "message": "Hash verified against FIR document"
            }

    if evidence_doc:
        return {
            "verified": True,
            "type": "evidence",
            "evidence_id": evidence_doc["evidence_id"],
            "filename": evidence_doc.get("original_filename"),
            "uploaded_at": evidence_doc["uploaded_at"],
            "uploaded_by": evidence_doc["uploaded_by"],
            "cid": evidence_doc.get("cid"),
            "message": "Hash verified against evidence"
        }

    return {
        "verified": False,
        "message": "Hash not found in system"
    }

@api.post("/verify/cid")
async def verify_cid(body: CIDVerificationIn, user: dict = Depends(require_roles("investigator", "forensic", "court_officer", "judge", "admin"))):
    # Search for CID in FIR documents and evidence
    fir_doc = await db.firs.find_one({"documents.cid": body.cid}, {"_id": 0})
    evidence_doc = await db.evidence.find_one({"cid": body.cid}, {"_id": 0})

    if fir_doc:
        doc = next((d for d in fir_doc.get("documents", []) if d.get("cid") == body.cid), None)
        if doc:
            return {
                "verified": True,
                "type": "fir_document",
                "fir_id": fir_doc["fir_id"],
                "filename": doc["filename"],
                "uploaded_at": doc["uploaded_at"],
                "uploaded_by": doc["uploaded_by"],
                "sha256_hash": doc["sha256_hash"],
                "message": "CID verified against FIR document"
            }

    if evidence_doc:
        return {
            "verified": True,
            "type": "evidence",
            "evidence_id": evidence_doc["evidence_id"],
            "filename": evidence_doc.get("original_filename"),
            "uploaded_at": evidence_doc["uploaded_at"],
            "uploaded_by": evidence_doc["uploaded_by"],
            "sha256_hash": evidence_doc["sha256_hash"],
            "message": "CID verified against evidence"
        }

    return {
        "verified": False,
        "message": "CID not found in system"
    }


@api.post("/ai/ask")
async def ai_ask(body: AIAskIn, user: dict = Depends(get_current_user)):
    if not body.question.strip():
        raise HTTPException(400, "Question is required")
    return _answer_ai_question(body.question, body.city or "", body.area or "")


@api.post("/smart-contract/validate")
async def smart_contract_validate(body: SmartContractValidationIn, user: dict = Depends(require_roles("judge", "court_officer", "investigator", "forensic", "police"))):
    payload = body.model_dump() if hasattr(body, "model_dump") else body.dict()
    logger.info(f"Smart contract validation payload received: {payload}")

    ev = await db.evidence.find_one({"evidence_id": body.evidence_id}, {"_id": 0})
    if not ev:
        raise HTTPException(404, "Evidence not found")

    candidate_hash = (body.hash_value or "").strip()
    if len(candidate_hash) != 64 or any(c not in "0123456789abcdefABCDEF" for c in candidate_hash):
        raise HTTPException(400, "Invalid hash")

    hash_match = ev.get("sha256_hash") == candidate_hash
    approvals_valid = body.approval_count >= 2
    accepted = hash_match and approvals_valid
    stored_approval_count = len(ev.get("approvals", []) or [])
    status = "Accepted" if accepted else "Rejected"
    message = (
        "Smart contract accepted the evidence." if accepted else
        "Smart contract rejected the evidence. Ensure the hash matches and approvals are at least 2."
    )

    return {
        "accepted": accepted,
        "status": status,
        "message": message,
        "hash_match": hash_match,
        "approvals_valid": approvals_valid,
        "stored_approval_count": stored_approval_count,
        "provided_approval_count": body.approval_count,
    }


@api.get("/analytics/stats")
async def analytics(user: dict = Depends(get_current_user)):
    def _aggregate_firs(firs_list):
        by_type = {}
        by_location = {}
        by_month = {}
        for f in firs_list:
            ct = f.get("crime_type") or "Unknown"
            loc = f.get("location") or "Unknown"
            by_type[ct] = by_type.get(ct, 0) + 1
            by_location[loc] = by_location.get(loc, 0) + 1
            created = f.get("created_at") or ""
            m = created[:7] if len(created) >= 7 else "unknown"
            by_month[m] = by_month.get(m, 0) + 1
        return by_type, by_location, by_month

    if user.get("role") == "citizen":
        cid = user["id"]
        firs = await db.firs.find({"citizen_id": cid}, {"_id": 0}).to_list(2000)
        my_fir_ids = [f["fir_id"] for f in firs if f.get("fir_id")]
        total = len(firs)
        closed = sum(1 for f in firs if is_case_closed_status(f.get("status")))
        open_cases = total - closed
        reopened_cases = sum(1 for f in firs if f.get("status") == "REOPENED")
        by_type, by_location, by_month = _aggregate_firs(firs)
        cases = await db.cases.find({"citizen_id": cid}, {"_id": 0, "case_id": 1}).to_list(500)
        case_ids = [c["case_id"] for c in cases if c.get("case_id")]
        total_evidence = (
            await db.evidence.count_documents({"case_id": {"$in": case_ids}}) if case_ids else 0
        )
        total_suspects = (
            await db.suspects.count_documents({"associated_cases": {"$in": case_ids}}) if case_ids else 0
        )
        tampered_evidence = (
            await db.evidence.count_documents({"case_id": {"$in": case_ids}, "tampered": True})
            if case_ids
            else 0
        )
        pending_appeals = (
            await db.appeals.count_documents({"fir_id": {"$in": my_fir_ids}, "status": "Pending"})
            if my_fir_ids
            else 0
        )
        total_appeals = (
            await db.appeals.count_documents({"fir_id": {"$in": my_fir_ids}}) if my_fir_ids else 0
        )
        top_suspects = (
            await db.suspects.find({"associated_cases": {"$in": case_ids}}, {"_id": 0})
            .sort("risk_level", -1)
            .limit(10)
            .to_list(10)
            if case_ids
            else []
        )
        # Sentence stats for citizen's cases
        active_sentences = (
            await db.suspects.count_documents({"associated_cases": {"$in": case_ids}, "sentence_status": "Active Sentence"})
            if case_ids else 0
        )
        completed_sentences = (
            await db.suspects.count_documents({"associated_cases": {"$in": case_ids}, "sentence_status": "Completed Sentence"})
            if case_ids else 0
        )
        released_prisoners = (
            await db.suspects.count_documents({"associated_cases": {"$in": case_ids}, "sentence_status": "Released"})
            if case_ids else 0
        )
        return {
            "total_firs": total,
            "open_cases": open_cases,
            "closed_cases": closed,
            "reopened_cases": reopened_cases,
            "total_evidence": total_evidence,
            "total_suspects": total_suspects,
            "tampered_evidence": tampered_evidence,
            "pending_appeals": pending_appeals,
            "total_appeals": total_appeals,
            "active_sentences": active_sentences,
            "completed_sentences": completed_sentences,
            "released_prisoners": released_prisoners,
            "by_crime_type": [{"name": k, "value": v} for k, v in by_type.items()],
            "by_location": [{"name": k, "value": v} for k, v in by_location.items()],
            "by_month": [{"name": k, "value": v} for k, v in sorted(by_month.items())],
            "top_suspects": top_suspects,
        }

    firs = await db.firs.find({}, {"_id": 0}).to_list(2000)
    total = len(firs)
    closed = sum(1 for f in firs if is_case_closed_status(f.get("status")))
    open_cases = total - closed
    reopened_cases = sum(1 for f in firs if f.get("status") == "REOPENED")
    by_type, by_location, by_month = _aggregate_firs(firs)
    top_suspects = await db.suspects.find({}, {"_id": 0}).sort("risk_level", -1).limit(10).to_list(10)
    pending_appeals = await db.appeals.count_documents({"status": "Pending"})
    total_appeals = await db.appeals.count_documents({})
    active_sentences = await db.suspects.count_documents({"sentence_status": "Active Sentence"})
    completed_sentences = await db.suspects.count_documents({"sentence_status": "Completed Sentence"})
    released_prisoners = await db.suspects.count_documents({"sentence_status": "Released"})
    return {
        "total_firs": total,
        "open_cases": open_cases,
        "closed_cases": closed,
        "reopened_cases": reopened_cases,
        "total_evidence": await db.evidence.count_documents({}),
        "total_suspects": await db.suspects.count_documents({}),
        "tampered_evidence": await db.evidence.count_documents({"tampered": True}),
        "pending_appeals": pending_appeals,
        "total_appeals": total_appeals,
        "active_sentences": active_sentences,
        "completed_sentences": completed_sentences,
        "released_prisoners": released_prisoners,
        "by_crime_type": [{"name": k, "value": v} for k, v in by_type.items()],
        "by_location": [{"name": k, "value": v} for k, v in by_location.items()],
        "by_month": [{"name": k, "value": v} for k, v in sorted(by_month.items())],
        "top_suspects": top_suspects,
    }

@api.get("/ai/predictions")
async def ai_predictions(user: dict = Depends(get_current_user)):
    try:
        import pandas as pd
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import seaborn as sns
    except Exception:
        raise HTTPException(500, "AI analytics dependencies are missing")

    firs = await db.firs.find({}, {"_id": 0}).to_list(5000)
    suspects = await db.suspects.find({}, {"_id": 0}).to_list(5000)

    if len(firs) == 0:
        return {
            "monthly_crime_count": [],
            "yearly_crime_trends": [],
            "most_common_crime_types": [],
            "top_dangerous_areas": [],
            "area_risk_scores": [],
            "repeat_offenders": [],
            "trend_summary": "No FIR data available for AI prediction.",
            "charts": {},
        }

    df = pd.DataFrame(firs)
    if "created_at" in df.columns:
        dt = pd.to_datetime(df["created_at"], errors="coerce", utc=True)
    elif "incident_date" in df.columns:
        dt = pd.to_datetime(df["incident_date"], errors="coerce", utc=True)
    else:
        dt = pd.Series([pd.NaT] * len(df))
    df["__dt"] = dt
    df["crime_type"] = df.get("crime_type", "Unknown").fillna("Unknown").astype(str)
    df["location"] = df.get("location", "Unknown").fillna("Unknown").astype(str)

    monthly = (
        df.dropna(subset=["__dt"])
        .assign(month=lambda x: x["__dt"].dt.to_period("M").astype(str))
        .groupby("month")
        .size()
        .reset_index(name="count")
        .sort_values("month")
    )
    yearly = (
        df.dropna(subset=["__dt"])
        .assign(year=lambda x: x["__dt"].dt.year.astype(str))
        .groupby("year")
        .size()
        .reset_index(name="count")
        .sort_values("year")
    )
    common_types = (
        df.groupby("crime_type")
        .size()
        .reset_index(name="count")
        .sort_values("count", ascending=False)
    )
    hotspots = (
        df.groupby("location")
        .size()
        .reset_index(name="count")
        .sort_values("count", ascending=False)
    )
    top_areas = hotspots.head(5).copy()

    # Rule-based area risk scoring from FIR density.
    def risk_from_count(count: int) -> str:
        if count >= 10:
            return "High"
        if count >= 5:
            return "Medium"
        return "Low"

    risk_scores = []
    for _, row in top_areas.iterrows():
        risk_scores.append({
            "area": row["location"],
            "count": int(row["count"]),
            "risk_level": risk_from_count(int(row["count"])),
        })

    repeat_offenders = []
    for s in suspects:
        linked = s.get("associated_cases") or []
        if len(linked) > 1:
            repeat_offenders.append({
                "suspect_id": s.get("suspect_id"),
                "name": s.get("name"),
                "linked_cases_count": len(linked),
                "linked_cases": linked,
            })
    repeat_offenders = sorted(repeat_offenders, key=lambda x: x["linked_cases_count"], reverse=True)[:20]

    sns.set_theme(style="whitegrid")
    ts = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    monthly_file = f"ai_monthly_trend_{ts}.png"
    yearly_file = f"ai_yearly_trend_{ts}.png"
    crime_type_file = f"ai_crime_types_{ts}.png"
    hotspot_file = f"ai_hotspots_{ts}.png"

    if len(monthly) > 0:
        plt.figure(figsize=(8, 4))
        sns.lineplot(data=monthly, x="month", y="count", marker="o")
        plt.xticks(rotation=30, ha="right")
        plt.title("Monthly Crime Trend")
        plt.tight_layout()
        plt.savefig(GRAPHS_DIR / monthly_file, dpi=140)
        plt.close()

    if len(yearly) > 0:
        plt.figure(figsize=(6, 4))
        sns.barplot(data=yearly, x="year", y="count", color="#0033A0")
        plt.title("Yearly Crime Trend")
        plt.tight_layout()
        plt.savefig(GRAPHS_DIR / yearly_file, dpi=140)
        plt.close()

    if len(common_types) > 0:
        top_types_chart = common_types.head(8)
        plt.figure(figsize=(8, 4))
        sns.barplot(data=top_types_chart, x="crime_type", y="count", color="#D92D20")
        plt.xticks(rotation=30, ha="right")
        plt.title("Most Common Crime Types")
        plt.tight_layout()
        plt.savefig(GRAPHS_DIR / crime_type_file, dpi=140)
        plt.close()

    if len(top_areas) > 0:
        plt.figure(figsize=(7, 4))
        plt.pie(top_areas["count"], labels=top_areas["location"], autopct="%1.1f%%")
        plt.title("Top 5 Dangerous Areas")
        plt.tight_layout()
        plt.savefig(GRAPHS_DIR / hotspot_file, dpi=140)
        plt.close()

    monthly_peak = monthly.sort_values("count", ascending=False).head(1)
    peak_text = ""
    if len(monthly_peak) > 0:
        peak_text = f"Peak month: {monthly_peak.iloc[0]['month']} ({int(monthly_peak.iloc[0]['count'])} FIRs). "
    area_text = ""
    if len(top_areas) > 0:
        area_text = f"Top hotspot: {top_areas.iloc[0]['location']} ({int(top_areas.iloc[0]['count'])} FIRs). "
    repeat_text = f"Repeat offenders detected: {len(repeat_offenders)}."

    return {
        "monthly_crime_count": [{"name": r["month"], "value": int(r["count"])} for _, r in monthly.iterrows()],
        "yearly_crime_trends": [{"name": r["year"], "value": int(r["count"])} for _, r in yearly.iterrows()],
        "most_common_crime_types": [{"name": r["crime_type"], "value": int(r["count"])} for _, r in common_types.head(10).iterrows()],
        "top_dangerous_areas": [{"name": r["location"], "value": int(r["count"])} for _, r in top_areas.iterrows()],
        "area_risk_scores": risk_scores,
        "repeat_offenders": repeat_offenders,
        "trend_summary": f"{peak_text}{area_text}{repeat_text}".strip(),
        "charts": {
            "monthly_trend": f"/graphs/{monthly_file}" if (GRAPHS_DIR / monthly_file).exists() else None,
            "yearly_trend": f"/graphs/{yearly_file}" if (GRAPHS_DIR / yearly_file).exists() else None,
            "crime_types_bar": f"/graphs/{crime_type_file}" if (GRAPHS_DIR / crime_type_file).exists() else None,
            "hotspots_pie": f"/graphs/{hotspot_file}" if (GRAPHS_DIR / hotspot_file).exists() else None,
        },
    }

# ---------------- Users (Admin) ----------------
@api.get("/users")
async def list_users(user: dict = Depends(require_roles("admin"))):
    return await db.users.find({}, {"_id": 0, "password_hash": 0}).sort("created_at", -1).to_list(500)

@api.patch("/users/{user_id}")
async def update_user(user_id: str, body: UserUpdate, user: dict = Depends(require_roles("admin"))):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if "role" in updates and updates["role"] not in ROLES:
        raise HTTPException(400, "Invalid role")
    await db.users.update_one({"id": user_id}, {"$set": updates})
    await log_activity(user["email"], "user_update", user_id)
    return await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})

@api.delete("/users/{user_id}")
async def delete_user(user_id: str, user: dict = Depends(require_roles("admin"))):
    await db.users.delete_one({"id": user_id})
    await log_activity(user["email"], "user_delete", user_id)
    return {"ok": True}

# ---------------- Officers list (for assigning) ----------------
@api.get("/officers")
async def list_officers(user: dict = Depends(require_roles("police", "forensic"))):
    return await db.users.find(
        {"role": {"$in": ["police", "forensic"]}},
        {"_id": 0, "password_hash": 0}
    ).to_list(200)

# ---------------- Report ----------------
@api.get("/reports/crime")
async def crime_report(user: dict = Depends(get_current_user)):
    total = await db.firs.count_documents({})
    closed = await db.firs.count_documents({"status": "Closed"})
    evidence_count = await db.evidence.count_documents({})
    suspects_count = await db.suspects.count_documents({})
    tampered = await db.evidence.count_documents({"tampered": True})
    text = (
        "=== AI CRIMINAL CASE MANAGEMENT SYSTEM ===\n"
        f"Generated: {now_iso()}\n"
        f"Generated By: {user['email']} ({user['role']})\n"
        "\n--- SUMMARY ---\n"
        f"Total FIRs: {total}\n"
        f"Closed Cases: {closed}\n"
        f"Open Cases: {total - closed}\n"
        f"Evidence Items: {evidence_count}\n"
        f"Suspects: {suspects_count}\n"
        f"Tampered Evidence Detected: {tampered}\n"
    )
    return {"report": text}

# ---------------- Fraud Detection & Trust Scoring ----------------
TRUST_TRACKED_ROLES = {"police", "investigator", "forensic", "judge"}

def _days_between_iso(start_iso: Optional[str], end_iso: Optional[str]) -> Optional[float]:
    if not start_iso or not end_iso:
        return None
    try:
        s = datetime.fromisoformat(start_iso.replace("Z", "+00:00"))
        e = datetime.fromisoformat(end_iso.replace("Z", "+00:00"))
        return max((e - s).total_seconds() / 86400.0, 0.0)
    except Exception:
        return None

async def _compute_user_trust(user_doc: dict) -> dict:
    uid = user_doc.get("id")
    email = user_doc.get("email")
    role = user_doc.get("role")

    reasons = []
    score = 50.0

    # + points: successful case closures handled by this user.
    closed_cases = await db.cases.count_documents({
        "status": "Closed",
        "$or": [
            {"created_by": email},
            {"assigned_investigator_email": email},
            {"closed_by": email},
            {"forwarded_to_forensic_by": email},
        ],
    })
    if closed_cases > 0:
        bonus = min(closed_cases * 3, 20)
        score += bonus
        reasons.append(f"+{bonus} for {closed_cases} successful case closures")

    # + points for verified evidence submissions.
    verified_evidence = await db.evidence.count_documents({
        "uploaded_by": email,
        "tampered": False,
    })
    if verified_evidence > 0:
        bonus = min(verified_evidence * 1.5, 15)
        score += bonus
        reasons.append(f"+{bonus:.1f} for verified evidence submissions")

    # +/- points for speed (fast vs delayed handling).
    owned_cases = await db.cases.find({
        "$or": [
            {"created_by": email},
            {"assigned_investigator_email": email},
            {"closed_by": email},
        ]
    }, {"_id": 0, "created_at": 1, "closed_at": 1, "status": 1}).to_list(500)
    durations = []
    for c in owned_cases:
        d = _days_between_iso(c.get("created_at"), c.get("closed_at"))
        if d is not None:
            durations.append(d)
    if durations:
        avg_days = sum(durations) / len(durations)
        if avg_days <= 15:
            score += 10
            reasons.append("+10 for fast case handling")
        elif avg_days > 45:
            score -= 10
            reasons.append("-10 for delayed case handling")

    # - points for complaints against user.
    complaints = int(user_doc.get("complaints") or 0)
    if complaints > 0:
        penalty = min(complaints * 5, 25)
        score -= penalty
        reasons.append(f"-{penalty} due to complaints ({complaints})")

    # - points for evidence mismatches / tampering uploaded by user.
    tampered_evidence = await db.evidence.count_documents({
        "uploaded_by": email,
        "tampered": True,
    })
    if tampered_evidence > 0:
        penalty = min(tampered_evidence * 12, 36)
        score -= penalty
        reasons.append(f"-{penalty} for evidence mismatch/tampering")

    # - points for rejected decisions (modeled by further investigation after judgment).
    rejected_decisions = await db.cases.count_documents({
        "status_history": {
            "$elemMatch": {
                "status": "Judgment Issued",
                "by": email,
            }
        },
        "judgment_verdict": "Further Investigation",
    })
    if rejected_decisions > 0:
        penalty = min(rejected_decisions * 8, 24)
        score -= penalty
        reasons.append(f"-{penalty} for rejected/returned decisions")

    score = max(0.0, min(100.0, round(score, 2)))
    if score < 20:
        risk = "High Risk"
        action = "Suspended"
    elif score < 40:
        risk = "Suspicious"
        action = "Watch List"
    else:
        risk = "Normal"
        action = "Normal"

    await db.trust_scores.update_one(
        {"user_id": uid},
        {"$set": {
            "user_id": uid,
            "email": email,
            "role": role,
            "trust_score": score,
            "risk_level": risk,
            "action": action,
            "reasons": reasons,
            "updated_at": now_iso(),
        }},
        upsert=True,
    )

    return {
        "user_id": uid,
        "name": user_doc.get("name"),
        "email": email,
        "role": role,
        "trust_score": score,
        "risk_level": risk,
        "action": action,
        "reasons": reasons,
    }

@api.get("/fraud/score/{user_id}")
async def fraud_score_user(user_id: str, user: dict = Depends(require_roles("admin"))):
    user_doc = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
    if not user_doc:
        raise HTTPException(404, "User not found")
    if user_doc.get("role") not in TRUST_TRACKED_ROLES:
        raise HTTPException(400, "Role not eligible for trust scoring")
    result = await _compute_user_trust(user_doc)
    await log_activity(user["email"], "fraud_score_user", user_id)
    return result

@api.get("/fraud/all-users")
async def fraud_all_users(user: dict = Depends(require_roles("admin"))):
    users = await db.users.find(
        {"role": {"$in": list(TRUST_TRACKED_ROLES)}},
        {"_id": 0, "password_hash": 0},
    ).to_list(1000)
    results = []
    for u in users:
        results.append(await _compute_user_trust(u))
    results.sort(key=lambda x: x["trust_score"])
    return {"count": len(results), "users": results}

@api.get("/fraud/flags")
async def fraud_flags(user: dict = Depends(require_roles("admin"))):
    users = await db.users.find(
        {"role": {"$in": list(TRUST_TRACKED_ROLES)}},
        {"_id": 0, "password_hash": 0},
    ).to_list(1000)
    flagged = []
    for u in users:
        scored = await _compute_user_trust(u)
        if scored["risk_level"] in ("Suspicious", "High Risk"):
            flagged.append(scored)
    flagged.sort(key=lambda x: x["trust_score"])
    await log_activity(user["email"], "fraud_flags_view", f"count={len(flagged)}")
    return {"count": len(flagged), "flags": flagged}


def _is_user_untrusted(user_doc: dict) -> bool:
    trust_score = float(user_doc.get("trust_score") or 0)
    status = str(user_doc.get("status") or "").strip().lower()
    flagged = bool(user_doc.get("flagged")) or status == "flagged"
    return trust_score < 50 or flagged


async def select_best_user(role: str, excluded_user_id: Optional[str] = None) -> Optional[dict]:
    query = {
        "role": role,
        "status": "active",
        "trust_score": {"$gte": 50},
    }
    if excluded_user_id:
        query["id"] = {"$ne": excluded_user_id}
    candidates = await db.users.find(
        query,
        {"_id": 0, "password_hash": 0},
    ).sort([
        ("trust_score", -1),
        ("complaints", 1),
        ("created_at", 1),
    ]).to_list(200)
    return candidates[0] if candidates else None


def _get_case_assigned_reference(case_doc: dict, role: str) -> dict:
    role = (role or "").strip().lower()
    assigned_users = case_doc.get("assigned_users")
    if isinstance(assigned_users, dict) and isinstance(assigned_users.get(role), dict):
        role_entry = assigned_users.get(role) or {}
        return {
            "id": role_entry.get("id") or role_entry.get("user_id"),
            "email": role_entry.get("email"),
        }

    role_id_key = f"assigned_{role}_id"
    role_email_key = f"assigned_{role}_email"
    if case_doc.get(role_id_key) or case_doc.get(role_email_key):
        return {
            "id": case_doc.get(role_id_key),
            "email": case_doc.get(role_email_key),
        }

    if role == "investigator":
        return {
            "id": case_doc.get("assigned_investigator_id"),
            "email": case_doc.get("assigned_investigator_email"),
        }
    return {"id": None, "email": None}


async def _find_user_by_ref(user_ref: dict) -> Optional[dict]:
    if user_ref.get("id"):
        user_doc = await db.users.find_one({"id": user_ref["id"]}, {"_id": 0, "password_hash": 0})
        if user_doc:
            return user_doc
    if user_ref.get("email"):
        return await db.users.find_one({"email": user_ref["email"]}, {"_id": 0, "password_hash": 0})
    return None


def _build_case_reassignment_updates(case_doc: dict, role: str, replacement: dict) -> dict:
    role = (role or "").strip().lower()
    updates = {
        "updated_at": now_iso(),
    }

    assigned_users = case_doc.get("assigned_users")
    if isinstance(assigned_users, dict):
        new_assigned_users = dict(assigned_users)
        new_assigned_users[role] = {
            "id": replacement["id"],
            "email": replacement.get("email", ""),
            "name": replacement.get("name", ""),
        }
        updates["assigned_users"] = new_assigned_users

    role_id_key = f"assigned_{role}_id"
    role_email_key = f"assigned_{role}_email"
    role_name_key = f"assigned_{role}_name"
    if role_id_key in case_doc or role_email_key in case_doc or role_name_key in case_doc:
        updates[role_id_key] = replacement["id"]
        updates[role_email_key] = replacement.get("email", "")
        updates[role_name_key] = replacement.get("name", "")

    if role == "investigator" or any(k in case_doc for k in ("assigned_investigator_id", "assigned_investigator_email", "assigned_investigator_name")):
        updates["assigned_investigator_id"] = replacement["id"]
        updates["assigned_investigator_email"] = replacement.get("email", "")
        updates["assigned_investigator_name"] = replacement.get("name", "")

    return updates


@api.post("/consensus/replace-user")
async def consensus_replace_user(body: ConsensusReplaceUserIn, user: dict = Depends(require_roles("admin"))):
    role = (body.role or "").strip().lower()
    if role not in ROLES:
        raise HTTPException(400, "Invalid role")

    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")

    assigned_ref = _get_case_assigned_reference(case_doc, role)
    removed_user = await _find_user_by_ref(assigned_ref)
    if not removed_user:
        raise HTTPException(404, "Assigned user not found for this role")

    if not _is_user_untrusted(removed_user):
        raise HTTPException(400, "Assigned user is not untrusted")

    await db.users.update_one(
        {"id": removed_user["id"]},
        {"$set": {"status": "untrusted"}},
    )

    replacement = await select_best_user(role, excluded_user_id=removed_user["id"])
    if not replacement:
        raise HTTPException(404, "No active trusted replacement user available")

    case_updates = _build_case_reassignment_updates(case_doc, role, replacement)
    event = {
        "type": "consensus_user_replaced",
        "role": role,
        "removed_user_id": removed_user["id"],
        "removed_user_email": removed_user.get("email", ""),
        "new_user_id": replacement["id"],
        "new_user_email": replacement.get("email", ""),
        "at": now_iso(),
        "by": user["email"],
    }
    await db.cases.update_one(
        {"case_id": body.case_id},
        {"$set": case_updates, "$push": {"consensus_updates": event}},
    )
    await append_block(
        body.case_id,
        "CONSENSUS_USER_REPLACED",
        {
            "role": role,
            "removed_user_id": removed_user["id"],
            "new_user_id": replacement["id"],
        },
        user["email"],
    )
    await log_activity(
        user["email"],
        "consensus_replace_user",
        f"{body.case_id} {role}: {removed_user.get('email', removed_user['id'])} -> {replacement.get('email', replacement['id'])}",
    )

    return {
        "removed_user": removed_user.get("email", removed_user["id"]),
        "new_user": replacement.get("email", replacement["id"]),
        "message": "User replaced successfully",
    }


# ---------------- Admin Backend APIs ----------------
@api.get("/admin/users")
async def admin_list_users(user: dict = Depends(require_roles("admin"))):
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).sort("created_at", -1).to_list(1000)
    return {"count": len(users), "users": users}


@api.patch("/admin/user/role")
async def admin_update_user_role(body: AdminUserRoleUpdateIn, user: dict = Depends(require_roles("admin"))):
    if body.role not in ROLES:
        raise HTTPException(400, "Invalid role")
    target = await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})
    if not target:
        raise HTTPException(404, "User not found")
    await db.users.update_one({"id": body.user_id}, {"$set": {"role": body.role}})
    await log_activity(user["email"], "admin_user_role_update", f"{body.user_id} → {body.role}")
    return await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})


@api.patch("/admin/user/status")
async def admin_update_user_status(body: AdminUserStatusUpdateIn, user: dict = Depends(require_roles("admin"))):
    target = await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})
    if not target:
        raise HTTPException(404, "User not found")
    st = (body.status or "").strip().lower()
    if st not in ("active", "suspended"):
        raise HTTPException(400, "status must be 'active' or 'suspended'")
    await db.users.update_one({"id": body.user_id}, {"$set": {"status": st}})
    await log_activity(user["email"], "admin_user_status_update", f"{body.user_id} → {st}")
    return await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})


@api.patch("/admin/user/queue-role")
async def admin_queue_role_request(body: AdminQueueRoleIn, user: dict = Depends(require_roles("admin"))):
    if body.pending_role not in ROLES or body.pending_role in ("admin", "citizen"):
        raise HTTPException(400, "Invalid queued role")
    target = await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})
    if not target:
        raise HTTPException(404, "User not found")
    if target.get("role") == body.pending_role:
        raise HTTPException(400, "User already holds this role")
    await db.users.update_one(
        {"id": body.user_id},
        {"$set": {"pending_role": body.pending_role, "pending_role_queued_at": now_iso()}},
    )
    await log_activity(user["email"], "admin_role_request_queued", f"{body.user_id} → {body.pending_role}")
    return await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})


@api.post("/admin/user/approve-role-request")
async def admin_approve_role_request(body: AdminTargetUserIn, user: dict = Depends(require_roles("admin"))):
    target = await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})
    if not target:
        raise HTTPException(404, "User not found")
    pending = target.get("pending_role")
    if not pending or pending not in ROLES or pending in ("admin", "citizen"):
        raise HTTPException(400, "No valid pending role request for this user")
    await db.users.update_one(
        {"id": body.user_id},
        {"$set": {"role": pending}, "$unset": {"pending_role": "", "pending_role_queued_at": ""}},
    )
    await log_activity(user["email"], "admin_role_request_approved", f"{body.user_id} → {pending}")
    return await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})


@api.post("/admin/user/reject-role-request")
async def admin_reject_role_request(body: AdminTargetUserIn, user: dict = Depends(require_roles("admin"))):
    target = await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})
    if not target:
        raise HTTPException(404, "User not found")
    if not target.get("pending_role"):
        raise HTTPException(400, "No pending role request for this user")
    await db.users.update_one(
        {"id": body.user_id},
        {"$unset": {"pending_role": "", "pending_role_queued_at": ""}},
    )
    await log_activity(user["email"], "admin_role_request_rejected", body.user_id)
    return await db.users.find_one({"id": body.user_id}, {"_id": 0, "password_hash": 0})


@api.get("/admin/cases")
async def admin_list_cases(user: dict = Depends(require_roles("admin"))):
    cases = await db.cases.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return {"count": len(cases), "cases": cases}


@api.patch("/admin/case/assign")
async def admin_assign_case(body: AdminCaseAssignIn, user: dict = Depends(require_roles("police"))):
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    investigator = await db.users.find_one({"id": body.investigator_id}, {"_id": 0, "password_hash": 0})
    if not investigator or investigator.get("role") != "investigator":
        raise HTTPException(404, "Investigator not found")
    await db.cases.update_one(
        {"case_id": body.case_id},
        {"$set": {
            "assigned_investigator_id": investigator["id"],
            "assigned_investigator_name": investigator.get("name", ""),
            "assigned_investigator_email": investigator.get("email", ""),
            "updated_at": now_iso(),
        }},
    )
    await log_activity(user["email"], "admin_case_assign", f"{body.case_id} → {investigator.get('email', investigator['id'])}")
    return await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})


@api.patch("/admin/case/status")
async def admin_update_case_status(body: AdminCaseStatusUpdateIn, user: dict = Depends(require_roles("police"))):
    if body.status not in CASE_STATUSES:
        raise HTTPException(400, "Invalid status")
    case_doc = await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(404, "Case not found")
    await db.cases.update_one(
        {"case_id": body.case_id},
        {"$set": {"status": body.status, "updated_at": now_iso()}},
    )
    await log_activity(user["email"], "admin_case_status_update", f"{body.case_id} → {body.status}")
    return await db.cases.find_one({"case_id": body.case_id}, {"_id": 0})


@api.get("/admin/fraud")
async def admin_list_fraud(user: dict = Depends(require_roles("admin"))):
    trust_scores = await db.trust_scores.find({}, {"_id": 0}).sort("updated_at", -1).to_list(1000)
    flags = await db.fraud_flags.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return {
        "trust_score_count": len(trust_scores),
        "flag_count": len(flags),
        "trust_scores": trust_scores,
        "flags": flags,
    }


@api.get("/admin/blockchain")
async def admin_list_blockchain(user: dict = Depends(require_roles("admin"))):
    blocks = await db.blockchain.find({}, {"_id": 0}).sort("index", -1).to_list(1000)
    return {"count": len(blocks), "blocks": blocks}


@api.get("/admin/evidence")
async def admin_list_evidence(user: dict = Depends(require_roles("admin"))):
    evidence = await db.evidence.find({}, {"_id": 0}).sort("uploaded_at", -1).to_list(1000)
    return {"count": len(evidence), "evidence": evidence}


@api.get("/admin/logs")
async def admin_list_logs(user: dict = Depends(require_roles("admin"))):
    logs = await db.activity_logs.find({}, {"_id": 0}).sort("timestamp", -1).limit(1000).to_list(1000)
    return {"count": len(logs), "logs": logs}

# ---------------- Startup ----------------
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    await db.users.create_index("email", unique=True)
    await db.users.create_index("id", unique=True)
    await db.firs.create_index("fir_id", unique=True)
    await db.cases.create_index("case_id", unique=True)
    await db.cases.create_index("fir_id", unique=True)
    await db.investigation_notes.create_index("case_id")
    await db.blockchain.create_index("index")
    await db.blockchain.create_index("current_hash", unique=True)
    admin_email = os.environ.get("ADMIN_EMAIL", "admin@cjs.gov")
    admin_password = os.environ.get("ADMIN_PASSWORD", "Admin@123")
    existing = await db.users.find_one({"email": admin_email})
    if existing is None:
        await db.users.insert_one({
            "id": str(uuid.uuid4()),
            "wallet_id": generate_wallet_id(),
            "email": admin_email,
            "password_hash": hash_password(admin_password),
            "name": "System Admin",
            "cnic": "",
            "phone": "",
            "role": "admin",
            "trust_score": 100.0,
            "complaints": 0,
            "created_at": now_iso(),
        })
    elif not verify_password(admin_password, existing["password_hash"]):
        await db.users.update_one({"email": admin_email}, {"$set": {"password_hash": hash_password(admin_password)}})
    
    yield
    
    # Shutdown
    client.close()

app = FastAPI(title="AI Criminal Case Management System", lifespan=lifespan)

# ---------------- Middleware & mount ----------------
app.include_router(api)
app.mount("/graphs", StaticFiles(directory=str(GRAPHS_DIR)), name="graphs")

# ---------------- Frontend Entrypoints ----------------
# Landing page at `/`.
@app.get("/", response_class=HTMLResponse)
async def landing():
    template_path = ROOT_DIR / "backend" / "landing_template.html"
    if template_path.exists():
        return template_path.read_text(encoding="utf-8")
    # Fallback if file location differs.
    fallback_path = ROOT_DIR / "landing_template.html"
    if fallback_path.exists():
        return fallback_path.read_text(encoding="utf-8")
    return HTMLResponse(
        "<html><body><h2>AI Criminal Case Management System</h2><p>Server is running.</p></body></html>"
    )

# SPA fallback: serve frontend for client-side routes (but never for API/auth routes).
@app.get("/{path:path}", response_class=HTMLResponse)
async def spa_fallback(path: str, request: Request):
    # Exclude API routes and auth routes.
    if path.startswith("api/"):
        raise HTTPException(404, "Not Found")
    if path.startswith("auth/"):
        raise HTTPException(404, "Not Found")
    if path == "login" or path == "register":
        raise HTTPException(404, "Not Found")
    if path.startswith("login") or path.startswith("register"):
        raise HTTPException(404, "Not Found")

    # Try serving the built React index.html from common locations.
    # If you use a production build, place it accordingly.
    candidates = [
        ROOT_DIR / "frontend" / "build" / "index.html",
        ROOT_DIR / "frontend" / "dist" / "index.html",
    ]
    for c in candidates:
        if c.exists():
            return c.read_text(encoding="utf-8")

    # If React build doesn't exist, return 404 rather than breaking API requests.
    raise HTTPException(404, "Not Found")

# CORS Configuration - Allow frontend origins
cors_origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://localhost:3000",
    "https://127.0.0.1:3000",
]
frontend_url = os.environ.get("FRONTEND_URL", "").strip()
if frontend_url and frontend_url not in cors_origins:
    cors_origins.append(frontend_url)

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO)
